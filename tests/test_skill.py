from __future__ import annotations

import copy
import json
import os
import re
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest import mock
from xml.etree import ElementTree as ET

import yaml
from docx import Document
from docx.shared import Pt

REPO = Path(__file__).resolve().parents[1]
SKILL = REPO / "skills" / "build-k12-research-project"
SCRIPTS = SKILL / "scripts"
EXAMPLE = SKILL / "references" / "project-manifest.example.json"
sys.path.insert(0, str(SCRIPTS))

import apply_docx_format_contract  # noqa: E402
import audit_content_integrity  # noqa: E402
import audit_docx_format  # noqa: E402
import audit_docx_style_contract  # noqa: E402
import audit_lifecycle_coverage  # noqa: E402
import audit_project_package  # noqa: E402
import audit_xlsx_style_contract  # noqa: E402
import build_delivery_archive  # noqa: E402
import build_material_generation_plan  # noqa: E402
import format_contracts  # noqa: E402
import generate_attention_items  # noqa: E402
import generate_topic_candidates  # noqa: E402
import initialize_project_package  # noqa: E402
import plan_incremental_refresh  # noqa: E402
import project_workflow  # noqa: E402
import record_manual_acceptance  # noqa: E402
import record_policy_snapshot  # noqa: E402
import refresh_package_controls  # noqa: E402
import register_material_file  # noqa: E402
import run_project_preflight  # noqa: E402
import select_topic_and_prepare_intake  # noqa: E402
import validate_project_manifest  # noqa: E402


class SkillTests(unittest.TestCase):
    def setUp(self) -> None:
        self.example = json.loads(EXAMPLE.read_text(encoding="utf-8"))

    def test_frontmatter_and_agent_metadata(self) -> None:
        text = (SKILL / "SKILL.md").read_text(encoding="utf-8")
        match = re.match(r"^---\n(.*?)\n---\n", text, re.S)
        self.assertIsNotNone(match)
        metadata = yaml.safe_load(match.group(1))
        self.assertEqual(metadata["name"], "build-k12-research-project")
        self.assertIn("贵州省", metadata["description"])
        interface = yaml.safe_load((SKILL / "agents" / "openai.yaml").read_text(encoding="utf-8"))["interface"]
        self.assertGreaterEqual(len(interface["short_description"]), 25)
        self.assertLessEqual(len(interface["short_description"]), 64)
        self.assertIn("$build-k12-research-project", interface["default_prompt"])

    def test_all_local_markdown_links_exist(self) -> None:
        broken: list[tuple[str, str]] = []
        pattern = re.compile(r"\[[^\]]+\]\(([^)]+)\)")
        for source in SKILL.rglob("*.md"):
            for target in pattern.findall(source.read_text(encoding="utf-8")):
                if "://" in target or target.startswith("#"):
                    continue
                relative = target.split("#", 1)[0]
                if relative and not (source.parent / relative).exists():
                    broken.append((str(source.relative_to(SKILL)), target))
        self.assertEqual(broken, [])

    def test_no_machine_paths_or_private_values(self) -> None:
        forbidden_text = ("/Users/", "file://", "vscode://")
        phone = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")
        identity = re.compile(r"(?<!\d)\d{17}[0-9Xx](?!\d)")
        findings: list[str] = []
        for source in SKILL.rglob("*"):
            if not source.is_file() or "__pycache__" in source.parts:
                continue
            chunks: list[str] = []
            if source.suffix.lower() in {".md", ".py", ".mjs", ".yaml", ".json"}:
                chunks.append(source.read_text(encoding="utf-8", errors="ignore"))
            elif source.suffix.lower() in {".docx", ".xlsx"}:
                with zipfile.ZipFile(source) as archive:
                    for name in archive.namelist():
                        if name.endswith(("document.xml", "sharedStrings.xml", "core.xml", "custom.xml")) or (
                            name.startswith("xl/worksheets/") and name.endswith(".xml")
                        ):
                            chunks.append(archive.read(name).decode("utf-8", errors="ignore"))
            combined = "\n".join(chunks)
            if any(value in combined for value in forbidden_text) or phone.search(combined) or identity.search(combined):
                findings.append(str(source.relative_to(SKILL)))
        self.assertEqual(findings, [])

    def test_example_manifest_and_status_enums(self) -> None:
        errors, warnings = validate_project_manifest.validate(self.example)
        self.assertEqual(errors, [])
        self.assertEqual(warnings, [])

        bad_instrument = copy.deepcopy(self.example)
        bad_instrument["instruments"][0]["status"] = "unknown"
        errors, _ = validate_project_manifest.validate(bad_instrument)
        self.assertTrue(any("工具" in value and "status" in value for value in errors))

        bad_claim = copy.deepcopy(self.example)
        bad_claim["claims"][0]["status"] = "unknown"
        errors, _ = validate_project_manifest.validate(bad_claim)
        self.assertTrue(any("claims[1]" in value and "status" in value for value in errors))

    def test_custom_packages_still_require_attention_file(self) -> None:
        data = copy.deepcopy(self.example)
        data["materials"] = [item for item in data["materials"] if item["material_role"] != "attention-items"]
        errors, warnings = audit_lifecycle_coverage.audit(data, final=False)
        self.assertEqual(errors, [])
        self.assertTrue(any("attention-file" in value for value in warnings))

    def test_commitments_are_enforced_only_at_closing(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            manifest = root / "project-manifest.json"
            manifest.write_text(json.dumps(self.example, ensure_ascii=False), encoding="utf-8")
            errors, _ = audit_project_package.audit(manifest, root, final=True)
            self.assertFalse(any("承诺成果" in value and "尚未兑现" in value for value in errors))

            closing = copy.deepcopy(self.example)
            closing["generation_contract"].update(
                {"package_scope": "closing-kit", "truth_state": "closing", "delivery_state": "closing-ready"}
            )
            manifest.write_text(json.dumps(closing, ensure_ascii=False), encoding="utf-8")
            errors, _ = audit_project_package.audit(manifest, root, final=True)
            self.assertTrue(any("承诺成果" in value and "尚未兑现" in value for value in errors))

    def test_relative_evidence_paths_resolve_from_package_root(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            evidence_dir = root / "evidence"
            evidence_dir.mkdir()
            (evidence_dir / "raw.json").write_text("{}", encoding="utf-8")
            data = copy.deepcopy(self.example)
            data["evidence"][0].update(
                {
                    "status": "collected",
                    "delivery_included": True,
                    "source_file": "evidence/raw.json",
                    "collected_date": "2026-08-12",
                }
            )
            manifest = root / "project-manifest.json"
            manifest.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            errors, _ = audit_project_package.audit(manifest, root, final=True)
            self.assertFalse(any("source_file不存在" in value for value in errors))

            data["evidence"][0]["source_file"] = "evidence/missing.json"
            manifest.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            errors, _ = audit_project_package.audit(manifest, root, final=True)
            self.assertTrue(any("source_file不存在" in value for value in errors))

    def test_confidential_evidence_can_remain_in_controlled_custody(self) -> None:
        data = copy.deepcopy(self.example)
        data["evidence"][0].update(
            {
                "status": "collected",
                "delivery_included": False,
                "source_file": None,
                "collected_date": "2026-08-12",
                "custody_record": {
                    "owner": "学校课题档案管理员",
                    "locator": "校内受控档案柜A-01",
                    "verified_at": "2026-08-12",
                },
            }
        )
        errors, _ = validate_project_manifest.validate(data)
        self.assertFalse(any("证据 E-STU-SURVEY-RAW" in value for value in errors), errors)

    def test_delivery_and_truth_state_must_match(self) -> None:
        data = copy.deepcopy(self.example)
        data["generation_contract"].update({"delivery_state": "closing-ready", "truth_state": "planning"})
        errors, _ = validate_project_manifest.validate(data)
        self.assertTrue(any("closing-ready" in value and "truth_state" in value for value in errors))

    def test_attention_items_use_manifest_case_field_names(self) -> None:
        from datetime import date

        data = copy.deepcopy(self.example)
        data["cases"][0].update(
            {
                "status": "implemented",
                "implementation": {
                    "date": "2026-08-01",
                    "school": "示例学校",
                    "grade_class": "八年级1班",
                    "teacher_ids": ["P01"],
                    "participant_n": 40,
                    "actual_periods": 2,
                    "material_version": "1.0",
                    "deviations": ["无偏离"],
                    "data_cutoff": "2026-08-01",
                },
                "evidence_ids": ["E-STU-SURVEY-RAW"],
            }
        )
        issues = generate_attention_items.gather_issues(data, date(2026, 8, 12))
        self.assertFalse(any(issue.category == "案例真实性" for issue in issues), issues)

    def test_structured_bracket_placeholders_are_detected(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "draft.md"
            source.write_text("课题名称：【课题规范题目】\n负责人：【填写】", encoding="utf-8")
            errors, warnings = audit_content_integrity.audit(source, "submission", True, None, [])
            self.assertEqual(errors, [])
            self.assertTrue(any("占位内容" in value for value in warnings))

    def test_final_preflight_treats_warnings_as_blocking(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            manifest = root / "project-manifest.json"
            manifest.write_text(json.dumps(self.example, ensure_ascii=False), encoding="utf-8")
            with (
                mock.patch.object(run_project_preflight, "audit_package", return_value=([], ["test warning"])),
                mock.patch.object(run_project_preflight, "audit_lifecycle", return_value=([], [])),
            ):
                working = run_project_preflight.run(manifest, root, final=False)
                final = run_project_preflight.run(manifest, root, final=True)
            self.assertEqual(working["status"], "passed")
            self.assertEqual(final["status"], "failed")
            self.assertEqual(final["schema_version"], "1.4")
            self.assertEqual({item["id"] for item in final["manual_gates"]}, set(record_manual_acceptance.MANUAL_GATES))

    def test_initializer_creates_full_lifecycle_scaffold(self) -> None:
        intake_path = SKILL / "references" / "project-intake.example.json"
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory) / "示例材料包"
            manifest_path = initialize_project_package.initialize(intake_path, root, SKILL)
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(data["schema_version"], "1.6")
            self.assertEqual(data["manual_acceptance"]["status"], "pending")
            self.assertEqual(len(data["materials"]), 26)
            self.assertTrue(all(item["included_in_batch"] for item in data["materials"]))
            self.assertEqual(data["materials"][0]["status"], "draft")
            self.assertTrue((root / "00_课题材料包注意事项_真实性与待办清单_v0.1.docx").is_file())
            self.assertTrue((root / "M25_整套材料目录与交付索引_v0.1.docx").is_file())
            self.assertTrue((root / "05原始数据" / "M09_课题研究数据工作簿_v0.1.xlsx").is_file())
            errors, _ = validate_project_manifest.validate(data)
            self.assertEqual(errors, [])
            errors, warnings = audit_lifecycle_coverage.audit(data, final=False)
            self.assertEqual(errors, [])
            self.assertFalse(any("缺少生命周期组" in value for value in warnings), warnings)
            for material_id, filename in (
                ("M00", "00_课题材料包注意事项_真实性与待办清单_v0.1.docx"),
                ("M25", "M25_整套材料目录与交付索引_v0.1.docx"),
            ):
                style_errors, _ = audit_docx_style_contract.audit(root / filename, material_id)
                self.assertEqual(style_errors, [], material_id)

    def test_every_material_has_a_fixed_resolvable_format_contract(self) -> None:
        catalog_errors = format_contracts.validate_contract_catalog()
        self.assertEqual(catalog_errors, [])
        catalog = format_contracts.load_contracts()
        self.assertEqual(set(catalog["materials"]), {f"M{index:02d}" for index in range(26)})
        body = format_contracts.material_contract("M21")["roles"]["body"]
        self.assertEqual(
            (body["font_east_asia"], body["font_ascii"], body["size_pt"], body["first_line_indent_pt"]),
            ("仿宋_GB2312", "Times New Roman", 12.0, 24.0),
        )
        self.assertEqual(format_contracts.material_contract("M13")["roles"]["title"]["size_pt"], 16.0)

    def test_format_contract_apply_and_audit_detects_typography_drift(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "M21.docx"
            apply_docx_format_contract.apply(
                SKILL / "assets" / "templates" / "analysis-report.docx", target, "M21"
            )
            errors, _ = audit_docx_style_contract.audit(target, "M21")
            self.assertEqual(errors, [])

            document = Document(target)
            paragraph = next(value for value in document.paragraphs if value.text.strip())
            next(value for value in paragraph.runs if value.text.strip()).font.size = Pt(10)
            document.save(target)
            errors, _ = audit_docx_style_contract.audit(target, "M21")
            self.assertTrue(any("字号应为" in value for value in errors), errors)

    def test_role_audit_detects_heading_body_misclassification(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "source.docx"
            target = Path(directory) / "M11.docx"
            document = Document()
            document.add_paragraph("现状诊断与综合分析报告")
            document.add_paragraph("课题负责人：示例教师")
            document.add_paragraph("研究背景")
            document.add_paragraph("本研究立足课堂中已经观察到的真实问题，分析其表现、原因和改进条件。")
            document.save(source)
            result = apply_docx_format_contract.apply(source, target, "M11")
            self.assertEqual(result["roles_applied"]["cover_metadata"], 1)
            self.assertEqual(result["roles_applied"]["heading1"], 1)
            errors, _ = audit_docx_style_contract.audit(target, "M11")
            self.assertEqual(errors, [])

            drifted = Document(target)
            apply_docx_format_contract.apply_paragraph(
                drifted.paragraphs[2], "body", format_contracts.material_contract("M11")
            )
            drifted.save(target)
            errors, _ = audit_docx_style_contract.audit(target, "M11")
            self.assertTrue(any("标题/正文角色冲突" in value for value in errors), errors)

    def test_verified_policy_snapshot_must_be_project_fresh(self) -> None:
        intake = json.loads((SKILL / "references" / "project-intake.example.json").read_text(encoding="utf-8"))
        manifest = initialize_project_package.build_manifest(intake)
        manifest["sources"] = [
            {
                "id": "POLICY-2026-01", "title": "2026年度官方通知", "source_type": "official-notice",
                "verification_status": "verified", "verified_at": "2026-08-12", "valid_for_year": 2026,
                "locator": "https://example.gov.cn/official-notice", "used_in": ["M01"],
            },
            {
                "id": "TEMPLATE-2026-01", "title": "2026年度官方模板", "source_type": "official-template",
                "verification_status": "verified", "verified_at": "2026-08-12", "valid_for_year": 2026,
                "locator": "https://example.gov.cn/official-template", "used_in": ["M01"],
                "local_file": "01政策与立项/2026年度官方模板.docx", "retrieved_at": "2026-08-12",
                "source_sha256": "b" * 64,
            },
        ]
        manifest["submission_requirements"].update(
            {
                "status": "verified", "verified_at": "2026-08-12", "search_run_id": "policy-20260812-01",
                "searched_at": "2026-08-12", "official_portals_checked": ["贵州省教育厅"],
                "search_queries": ["2026 贵州省教育科学规划课题 申报 官方"],
                "policy_snapshot_file": "01政策与立项/当年要求核验快照_2026_2026-08-12.json",
                "policy_snapshot_sha256": "a" * 64, "notice_source_ids": ["POLICY-2026-01"],
                "template_source_ids": ["TEMPLATE-2026-01"], "submission_mode": "mixed",
            }
        )
        errors, _ = validate_project_manifest.validate(manifest)
        self.assertEqual(errors, [])
        manifest["submission_requirements"]["searched_at"] = "2026-08-04"
        errors, _ = validate_project_manifest.validate(manifest)
        self.assertTrue(any("超过7天" in value for value in errors), errors)

    def test_application_ready_is_blocked_after_deadline(self) -> None:
        data = copy.deepcopy(self.example)
        data["generation_contract"].update(
            {"package_scope": "application-kit", "delivery_state": "application-ready", "truth_state": "planning"}
        )
        data["submission_requirements"].update(
            {"status": "verified", "deadline": "2026-08-01", "searched_at": "2026-08-12"}
        )
        errors, _ = validate_project_manifest.validate(data)
        self.assertTrue(any("申报截止日" in value and "application-ready" in value for value in errors), errors)

    def test_record_policy_snapshot_writes_hashed_project_file(self) -> None:
        intake = SKILL / "references" / "project-intake.example.json"
        policy_example = json.loads(
            (SKILL / "references" / "policy-search-input.example.json").read_text(encoding="utf-8")
        )
        policy_example["search_run_id"] = "policy-test-20260812"
        policy_example["searched_at"] = "2026-08-12"
        for source in policy_example["sources"]:
            source["verified_at"] = "2026-08-12"
        with tempfile.TemporaryDirectory() as directory:
            base = Path(directory)
            root = base / "package"
            manifest_path = initialize_project_package.initialize(intake, root, SKILL)
            template_source = SKILL / "assets" / "templates" / "research-form.docx"
            for index, source in enumerate(
                [item for item in policy_example["sources"] if item["id"] in policy_example["template_source_ids"]],
                1,
            ):
                target = root / "01政策与立项" / f"official-template-{index}.docx"
                target.write_bytes(template_source.read_bytes())
                source.update(
                    local_file=str(target.relative_to(root)),
                    source_sha256=record_policy_snapshot.digest(target),
                    retrieved_at="2026-08-12",
                )
            policy_input = base / "policy.json"
            policy_input.write_text(json.dumps(policy_example, ensure_ascii=False), encoding="utf-8")
            snapshot = record_policy_snapshot.record(manifest_path, root, policy_input)
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            requirements = manifest["submission_requirements"]
            self.assertEqual(requirements["status"], "verified")
            self.assertEqual(requirements["policy_snapshot_file"], str(snapshot.relative_to(root.resolve())))
            self.assertEqual(requirements["policy_snapshot_sha256"], record_policy_snapshot.digest(snapshot))
            self.assertEqual(manifest["schema_version"], "1.6")
            self.assertTrue(manifest["materials"][1]["reference_template"])
            errors, _ = validate_project_manifest.validate(manifest)
            self.assertEqual(errors, [])

            manual_input = base / "manual.json"
            manual_input.write_text(
                json.dumps(
                    {
                        "reviewed_at": "2026-08-12T16:00:00+08:00",
                        "reviewer": "课题负责人",
                        "gates": {
                            gate_id: {"status": "passed", "note": f"已逐项复核：{description}"}
                            for gate_id, description in record_manual_acceptance.MANUAL_GATES.items()
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            attestation = record_manual_acceptance.record(manifest_path, root, manual_input)
            self.assertTrue(attestation.is_file())
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(manifest["manual_acceptance"]["status"], "verified")
            manifest["project"]["school"] = "变更后的示例学校"
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            final_errors, _ = audit_project_package.audit(manifest_path, root, final=True)
            self.assertTrue(any("主清单研究事实已变化" in value for value in final_errors), final_errors)
            refresh_package_controls.refresh(manifest_path)
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(manifest["manual_acceptance"]["status"], "expired")

    def test_scaffold_delivery_archive_contains_preflight_and_checksums(self) -> None:
        intake = SKILL / "references" / "project-intake.example.json"
        with tempfile.TemporaryDirectory() as directory:
            base = Path(directory)
            root = base / "package"
            manifest = initialize_project_package.initialize(intake, root, SKILL)
            output = base / "课题材料工作包.zip"
            result = build_delivery_archive.build(manifest, root, output, final=False)
            self.assertEqual(result["package_state"], "scaffold")
            with zipfile.ZipFile(output) as archive:
                self.assertIsNone(archive.testzip())
                names = set(archive.namelist())
                self.assertIn("project-manifest.json", names)
                self.assertIn("交付校验/preflight-report.json", names)
                self.assertIn("交付校验/SHA256SUMS.txt", names)

    def test_final_package_requires_manual_acceptance(self) -> None:
        intake = SKILL / "references" / "project-intake.example.json"
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory) / "package"
            manifest = initialize_project_package.initialize(intake, root, SKILL)
            errors, _ = audit_project_package.audit(manifest, root, final=True)
            self.assertTrue(any("人工终审确认" in value for value in errors), errors)

    def test_excluded_official_forms_do_not_create_scope_warnings(self) -> None:
        intake = json.loads((SKILL / "references" / "project-intake.example.json").read_text(encoding="utf-8"))
        intake["generation_contract"].update({"package_scope": "implementation-kit", "truth_state": "implementing"})
        manifest = initialize_project_package.build_manifest(intake)
        _, warnings = validate_project_manifest.validate(manifest)
        warning_text = "\n".join(warnings)
        self.assertNotIn("材料 M01", warning_text)
        self.assertNotIn("材料 M02", warning_text)
        self.assertNotIn("材料 M22", warning_text)
        self.assertIn("材料 M03", warning_text)

    def test_subject_coverage_matches_main_and_related_subjects(self) -> None:
        intake = json.loads((SKILL / "references" / "project-intake.example.json").read_text(encoding="utf-8"))
        intake["project"]["related_subjects"] = ["信息科技"]
        with self.assertRaisesRegex(ValueError, "subject_coverage学科集合"):
            initialize_project_package.build_manifest(intake)

        intake["subject_coverage"].append(
            {
                "subject": "信息科技",
                "role": "related",
                "research_function": "提供数字地图工具和电子作品载体",
                "standards_reference": "义务教育信息科技课程标准相应内容要求",
                "reviewer": "示例教师乙",
                "review_status": "confirmed",
            }
        )
        manifest = initialize_project_package.build_manifest(intake)
        errors, _ = validate_project_manifest.validate(manifest)
        self.assertEqual(errors, [])

    def test_major_k12_subject_domains_have_specific_profiles(self) -> None:
        subjects = (
            "语文", "数学", "英语", "物理", "化学", "生物", "小学科学", "道德与法治", "历史", "地理",
            "体育与健康", "音乐", "美术", "信息科技", "劳动", "综合实践", "心理健康", "幼小衔接",
            "特殊教育", "学校管理",
        )
        generic = generate_attention_items.GENERIC_SUBJECT_PROFILE
        self.assertEqual([value for value in subjects if generate_attention_items.profile_for_subject(value) is generic], [])

    def test_topic_generator_returns_5_to_8_scored_subject_aware_candidates(self) -> None:
        profile = json.loads((SKILL / "references" / "teacher-profile.example.json").read_text(encoding="utf-8"))
        for count in (5, 6, 7, 8):
            result = generate_topic_candidates.generate(profile, count)
            self.assertEqual(len(result["candidates"]), count)
            self.assertEqual(len({item["title"] for item in result["candidates"]}), count)
            self.assertEqual([item["priority_rank"] for item in result["candidates"]], list(range(1, count + 1)))
            for item in result["candidates"]:
                self.assertEqual(sum(item["score_explanation"].values()), item["score"])
                self.assertEqual(item["subject_coverage"][0]["subject"], "数学")
                self.assertEqual(item["subject_coverage"][0]["role"], "main")

    def test_topic_profile_validation_rejects_duplicate_subjects_and_invalid_counts(self) -> None:
        profile = json.loads((SKILL / "references" / "teacher-profile.example.json").read_text(encoding="utf-8"))
        profile["teaching"]["related_subjects"] = ["数学", "数学"]
        profile["teaching"]["student_count"] = True
        findings = generate_topic_candidates.validate_profile(profile)
        self.assertTrue(any("related_subjects" in value for value in findings), findings)
        self.assertTrue(any("student_count" in value for value in findings), findings)

    def test_local_resource_topic_requires_real_resource_context(self) -> None:
        profile = json.loads((SKILL / "references" / "teacher-profile.example.json").read_text(encoding="utf-8"))
        profile["school"]["context"] = "普通城区学校"
        profile["problem"]["available_resources"] = ["日常作业", "单元检测"]
        result = generate_topic_candidates.generate(profile, 5)
        self.assertFalse(any("乡土资源" in item["title"] for item in result["candidates"]))
        profile["problem"]["available_resources"].append("黔东南乡土数学建模资源")
        result = generate_topic_candidates.generate(profile, 5)
        self.assertTrue(any("乡土资源" in item["title"] for item in result["candidates"]))

    def test_cross_subject_candidates_register_each_real_subject(self) -> None:
        profile = json.loads((SKILL / "references" / "teacher-profile.example.json").read_text(encoding="utf-8"))
        profile["teaching"].update({"subject": "地理", "related_subjects": ["信息科技"]})
        profile["problem"]["description"] = "学生使用数字地图解释区域差异时只会找位置，不会组织证据"
        result = generate_topic_candidates.generate(profile, 8)
        cross = [item for item in result["candidates"] if item["route"].startswith("跨学科")]
        self.assertEqual(len(cross), 2)
        for item in cross:
            self.assertEqual([value["subject"] for value in item["subject_coverage"]], ["地理", "信息科技"])
            self.assertEqual([value["role"] for value in item["subject_coverage"]], ["main", "related"])

    def test_topic_profiles_cover_lab_art_pe_early_and_special_safety(self) -> None:
        profile = json.loads((SKILL / "references" / "teacher-profile.example.json").read_text(encoding="utf-8"))
        cases = {
            "化学": ("实验现象记录不完整，不能用证据解释反应", "危化品"),
            "体育与健康": ("学生在篮球练习中动作差异大且缺少过程反馈", "急救预案"),
            "美术": ("学生创作只重成品，缺少构思和修改过程", "图片版权"),
            "幼小衔接": ("幼儿入学准备活动被单次测评替代", "不过度测评"),
            "特殊教育": ("融合课堂中的个别化支持记录零散", "敏感信息保护"),
        }
        for subject, (problem, expected_risk) in cases.items():
            with self.subTest(subject=subject):
                profile["teaching"].update({"subject": subject, "related_subjects": []})
                profile["problem"]["description"] = problem
                result = generate_topic_candidates.generate(profile, 5)
                self.assertIn(expected_risk, "；".join(result["candidates"][0]["risks"]))
                self.assertTrue(result["candidates"][0]["evidence_plan"])

    def test_selected_topic_prepares_valid_intake_and_missing_fields_block(self) -> None:
        profile = json.loads((SKILL / "references" / "teacher-profile.example.json").read_text(encoding="utf-8"))
        topics = generate_topic_candidates.generate(profile, 6)
        selected = select_topic_and_prepare_intake.prepare(profile, topics, "TOPIC-02")
        self.assertTrue(selected["ready_for_initialization"])
        manifest = initialize_project_package.build_manifest(selected["project_intake"])
        errors, _ = validate_project_manifest.validate(manifest)
        self.assertEqual(errors, [])

        changed = select_topic_and_prepare_intake.prepare(profile, topics, "TOPIC-02", "真实问题导向的初中数学函数建模任务设计研究")
        self.assertTrue(changed["title_modified"])
        self.assertEqual(changed["project_intake"]["project"]["title"], "真实问题导向的初中数学函数建模任务设计研究")
        intake = selected["project_intake"]
        self.assertEqual(intake["project_context"]["grade_classes"], ["八年级1班", "八年级2班"])
        self.assertEqual(intake["problem_context"]["observed_evidence"], profile["problem"]["observed_evidence"])
        self.assertEqual([item["name"] for item in intake["commitments"]], ["最终研究报告", "教学案例集", "任务单", "评价量规"])

        with self.assertRaisesRegex(ValueError, "偏离所选候选方向"):
            select_topic_and_prepare_intake.prepare(profile, topics, "TOPIC-02", "小学生劳动习惯培养研究")

        changed_profile = copy.deepcopy(profile)
        changed_profile["problem"]["description"] = "完全不同的新问题"
        with self.assertRaisesRegex(ValueError, "重新生成候选题"):
            select_topic_and_prepare_intake.prepare(changed_profile, topics, "TOPIC-02")

        del profile["timeline"]["completion"]
        blocked = select_topic_and_prepare_intake.prepare(profile, topics, "TOPIC-02")
        self.assertFalse(blocked["ready_for_initialization"])
        self.assertIn("timeline.completion", blocked["missing_after_selection"])
        self.assertNotIn("project_intake", blocked)

    def test_generation_plan_has_every_material_truth_and_format_gate(self) -> None:
        intake = json.loads((SKILL / "references" / "project-intake.example.json").read_text(encoding="utf-8"))
        manifest = initialize_project_package.build_manifest(intake)
        plan = build_material_generation_plan.build(manifest)
        self.assertEqual(plan["material_job_count"], 26)
        self.assertEqual({job["material_id"] for job in plan["jobs"]}, {item["id"] for item in manifest["materials"]})
        application = next(job for job in plan["jobs"] if job["material_id"] == "M01")
        final_report = next(job for job in plan["jobs"] if job["material_id"] == "M21")
        self.assertEqual(application["execution_state"], "blocked-template")
        self.assertIn("当年官方模板或用户指定同类模板", application["truth_blockers_for_finalization"])
        self.assertIn("真实原始数据或实施记录", final_report["truth_blockers_for_finalization"])
        self.assertEqual(final_report["format_contract_id"], "final-report-fixed")
        self.assertEqual(final_report["resolved_format_contract"]["roles"]["body"]["first_line_indent_pt"], 24.0)
        self.assertTrue(all(job["source_manifest_fields"] and job["completion_gate"] for job in plan["jobs"]))
        self.assertTrue(plan["waiting_jobs"])

    def test_verified_material_is_not_treated_as_complete_and_waiting_is_explained(self) -> None:
        intake = json.loads((SKILL / "references" / "project-intake.example.json").read_text(encoding="utf-8"))
        manifest = initialize_project_package.build_manifest(intake)
        by_id = {item["id"]: item for item in manifest["materials"]}
        by_id["M00"].update(status="verified", file_path="attention.docx")
        plan = build_material_generation_plan.build(manifest)
        self.assertIn("JOB-M00", plan["next_jobs"])
        opening = next(job for job in plan["jobs"] if job["material_id"] == "M03")
        self.assertEqual(opening["execution_state"], "waiting-dependency")
        self.assertIn("M01", opening["waiting_for_material_ids"])

    def test_end_to_end_workflow_reaches_material_queue(self) -> None:
        profile_path = SKILL / "references" / "teacher-profile.example.json"
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory) / "workflow"
            start = project_workflow.start(profile_path, root, 6)
            self.assertEqual(start["state"], "awaiting-topic-selection")
            selected, code = project_workflow.select(root, "TOPIC-02")
            self.assertEqual(code, 0)
            self.assertEqual(selected["state"], "ready-to-initialize")
            state = project_workflow.initialize_workspace(root, SKILL)
            self.assertEqual(state["state"], "package-in-progress")
            plan = json.loads((root / "workflow-control" / "material-generation-plan.json").read_text(encoding="utf-8"))
            self.assertEqual(plan["material_job_count"], 26)
            self.assertIn("JOB-M00", plan["next_jobs"])
            self.assertIn("JOB-M01", plan["blocked_jobs"])
            self.assertNotIn("JOB-M01", plan["next_jobs"])
            self.assertIn("JOB-M03", plan["waiting_jobs"])

    def test_pending_truth_status_can_be_registered_and_controls_refreshed(self) -> None:
        intake_path = SKILL / "references" / "project-intake.example.json"
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory) / "package"
            manifest_path = initialize_project_package.initialize(intake_path, root, SKILL)
            data = json.loads(manifest_path.read_text(encoding="utf-8"))
            material = next(item for item in data["materials"] if item["id"] == "M11")
            target = root / material["planned_file_path"]
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes((SKILL / "assets" / "templates" / "analysis-report.docx").read_bytes())
            register_material_file.register(manifest_path, "M11", target, "pending-data", None, None)
            registered = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(next(item for item in registered["materials"] if item["id"] == "M11")["status"], "pending-data")
            for dependency_id in ("M09", "M10"):
                dependency = next(item for item in registered["materials"] if item["id"] == dependency_id)
                dependency.update(status="draft", file_path=dependency["planned_file_path"], sha256="0" * 64)
            plan = build_material_generation_plan.build(registered)
            job = next(item for item in plan["jobs"] if item["material_id"] == "M11")
            self.assertEqual(job["execution_state"], "waiting-truth-input")
            self.assertNotIn("JOB-M11", plan["next_jobs"])
            registered["samples"][0]["actual_n"] = registered["samples"][0]["planned_n"]
            registered["evidence"][0].update(
                status="collected", collected_date="2026-08-12", source_file="原始数据/问卷.xlsx",
                custody_record={"owner": "示例教师", "locator": "校内加密盘", "verified_at": "2026-08-12"},
            )
            self.assertTrue(build_material_generation_plan.truth_inputs_ready(registered, "pending-data"))
            result = refresh_package_controls.refresh(manifest_path)
            self.assertTrue(Path(result["attention"]).is_file())
            refreshed = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(next(item for item in refreshed["materials"] if item["id"] == "M00")["status"], "draft")

    def test_material_dependency_cycles_are_rejected(self) -> None:
        data = copy.deepcopy(self.example)
        data["materials"][0]["depends_on"] = ["M01"]
        data["materials"][1]["depends_on"] = ["M00"]
        errors, _ = validate_project_manifest.validate(data)
        self.assertTrue(any("材料依赖存在循环" in value for value in errors), errors)

    def test_draft_files_are_audited_by_preflight(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            draft = root / "draft.docx"
            draft.write_bytes((SKILL / "assets" / "templates" / "research-form.docx").read_bytes())
            data = copy.deepcopy(self.example)
            data["materials"][0].update(
                {
                    "status": "draft",
                    "file_path": draft.name,
                    "sha256": audit_project_package.sha256(draft),
                }
            )
            manifest = root / "project-manifest.json"
            manifest.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            with mock.patch.object(run_project_preflight, "audit_content", return_value=([], [])) as content_audit:
                run_project_preflight.run(manifest, root, final=False)
            content_audit.assert_called_once()

    def test_incremental_refresh_propagates_subject_change(self) -> None:
        old = copy.deepcopy(self.example)
        old["generation_contract"].update({"snapshot_id": "snap-1", "generated_at": "2026-08-12T00:00:00+08:00"})
        new = copy.deepcopy(old)
        new["generation_contract"].update(
            {"batch_mode": "incremental", "snapshot_id": "snap-2", "parent_snapshot_id": "snap-1"}
        )
        new["subject_coverage"] = [
            {
                "subject": "地理",
                "role": "main",
                "research_function": "区域认知",
                "standards_reference": "课程标准",
                "reviewer": "示例教师",
                "review_status": "confirmed",
            }
        ]
        report = plan_incremental_refresh.compare(old, new)
        self.assertEqual(report["errors"], [])
        roles = {item["material_role"] for item in report["affected_materials"]}
        self.assertIn("analysis-report", roles)
        self.assertIn("attention-items", roles)

    def test_ready_registration_requires_real_qa_report(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            data = copy.deepcopy(self.example)
            material = root / "attention.docx"
            material.write_bytes((SKILL / "assets" / "templates" / "attention-items.docx").read_bytes())
            manifest = root / "project-manifest.json"
            manifest.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "必须提供--qa-report"):
                register_material_file.register(manifest, "M00", material, "ready", None, None)

    def test_included_evidence_files_are_registered_and_hashed(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            evidence_file = root / "evidence" / "raw.json"
            evidence_file.parent.mkdir()
            evidence_file.write_text("{}", encoding="utf-8")
            data = copy.deepcopy(self.example)
            data["evidence"][0].update(
                {
                    "status": "collected",
                    "delivery_included": True,
                    "source_file": "evidence/raw.json",
                    "source_sha256": audit_project_package.sha256(evidence_file),
                    "collected_date": "2026-08-12",
                }
            )
            manifest = root / "project-manifest.json"
            manifest.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            errors, warnings = audit_project_package.audit(manifest, root, final=False)
            self.assertEqual(errors, [])
            self.assertFalse(any("未登记材料" in value and "raw.json" in value for value in warnings), warnings)

    def test_docx_templates_are_valid_and_structurally_audited(self) -> None:
        profiles = {
            "research-form.docx": "research-form",
            "analysis-report.docx": "analysis-report",
            "lesson-table.docx": "lesson-table",
            "lesson-long.docx": "lesson-long",
            "casebook.docx": "casebook",
            "evidence-sheet.docx": "evidence-sheet",
            "attention-items.docx": "attention-items",
        }
        templates = SKILL / "assets" / "templates"
        for filename, profile in profiles.items():
            source = templates / filename
            with zipfile.ZipFile(source) as archive:
                self.assertIsNone(archive.testzip(), filename)
                core = archive.read("docProps/core.xml").decode("utf-8", errors="ignore")
                self.assertNotRegex(core, r"<(?:dc:creator|cp:lastModifiedBy)>\s*[^<\s]")
            result = subprocess.run(
                [sys.executable, str(SCRIPTS / "audit_docx_format.py"), str(source), "--profile", profile],
                text=True,
                capture_output=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def test_docx_template_builder_is_reproducible(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            first = Path(directory) / "first"
            second = Path(directory) / "second"
            for target in (first, second):
                result = subprocess.run(
                    [sys.executable, str(SCRIPTS / "build_generic_docx_templates.py"), str(target)],
                    text=True, capture_output=True, check=False,
                )
                self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertEqual(
                {path.name: path.read_bytes() for path in first.glob("*.docx")},
                {path.name: path.read_bytes() for path in second.glob("*.docx")},
            )

    def test_official_exact_detects_table_geometry_change(self) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        reference = SKILL / "assets" / "templates" / "research-form.docx"
        with tempfile.TemporaryDirectory() as directory:
            changed = Path(directory) / "changed.docx"
            document = Document(reference)
            first_col = document.tables[0]._tbl.tblGrid.findall(qn("w:gridCol"))[0]
            first_col.set(qn("w:w"), str(int(first_col.get(qn("w:w"))) + 120))
            document.save(changed)
            errors, _ = audit_docx_format.audit(changed, "official-exact", reference, False)
            self.assertTrue(any("列宽网格" in value for value in errors), errors)

    def test_official_exact_can_allow_only_added_drawings(self) -> None:
        from docx import Document
        from docx.shared import Inches
        from PIL import Image

        reference = SKILL / "assets" / "templates" / "research-form.docx"
        with tempfile.TemporaryDirectory() as directory:
            directory_path = Path(directory)
            image_path = directory_path / "signature.png"
            Image.new("RGB", (20, 20), "white").save(image_path)
            changed = directory_path / "changed.docx"
            document = Document(reference)
            document.paragraphs[0].add_run().add_picture(str(image_path), width=Inches(0.1))
            document.save(changed)
            errors, _ = audit_docx_format.audit(changed, "official-exact", reference, False)
            self.assertTrue(any("绘图或图片数量" in value for value in errors), errors)
            errors, _ = audit_docx_format.audit(changed, "official-exact", reference, False, True)
            self.assertFalse(any("绘图或图片" in value for value in errors), errors)

    def test_xlsx_template_has_required_sheets_and_clean_formulas(self) -> None:
        source = SKILL / "assets" / "templates" / "project-data-workbook.xlsx"
        required = {"项目说明", "变量编码", "原始数据", "清理编码", "统计分析", "图表结果", "证据索引", "照片登记", "材料进度"}
        with zipfile.ZipFile(source) as archive:
            self.assertIsNone(archive.testzip())
            workbook = ET.fromstring(archive.read("xl/workbook.xml"))
            names = {element.attrib["name"] for element in workbook.iter() if element.tag.endswith("sheet")}
            formulas = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.startswith("xl/worksheets/") and name.endswith(".xml")
            )
            worksheet_roots = [
                ET.fromstring(archive.read(name))
                for name in archive.namelist()
                if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
            ]
        self.assertEqual(names, required)
        namespace = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
        self.assertEqual(sum(len(root.findall(f".//{namespace}pane")) for root in worksheet_roots), 9)
        self.assertGreaterEqual(sum(len(root.findall(f".//{namespace}dataValidation")) for root in worksheet_roots), 9)
        for token in ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "AVERIF"):
            self.assertNotIn(token, formulas)

        result = subprocess.run(
            [sys.executable, str(SCRIPTS / "audit_xlsx_structure.py"), str(source)],
            text=True,
            capture_output=True,
            check=False,
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        style_errors, style_warnings = audit_xlsx_style_contract.audit(source)
        self.assertEqual(style_errors, [])
        self.assertEqual(style_warnings, [])

        with tempfile.TemporaryDirectory() as directory:
            drifted = Path(directory) / "drifted.xlsx"
            with zipfile.ZipFile(source) as source_zip, zipfile.ZipFile(drifted, "w") as target_zip:
                for info in source_zip.infolist():
                    content = source_zip.read(info.filename)
                    if info.filename == "xl/styles.xml":
                        content = content.replace(b"Microsoft YaHei", b"Arial", 1)
                    target_zip.writestr(info, content)
            drift_errors, _ = audit_xlsx_style_contract.audit(drifted)
            self.assertTrue(any("styles.xml" in value or "全簿字体" in value for value in drift_errors), drift_errors)

    def test_final_standalone_audit_blocks_warnings(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "draft.md"
            source.write_text("【填写】", encoding="utf-8")
            result = subprocess.run(
                [sys.executable, str(SCRIPTS / "audit_content_integrity.py"), str(source), "--final"],
                text=True,
                capture_output=True,
                check=False,
            )
            self.assertEqual(result.returncode, 1, result.stdout + result.stderr)

    def test_scripts_compile_and_are_executable(self) -> None:
        scripts = [source for source in SCRIPTS.iterdir() if source.suffix in {".py", ".mjs"}]
        self.assertTrue(scripts)
        for source in scripts:
            self.assertTrue(os.access(source, os.X_OK), source.name)


if __name__ == "__main__":
    unittest.main()
