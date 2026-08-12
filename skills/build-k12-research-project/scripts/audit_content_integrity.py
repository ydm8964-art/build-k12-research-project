#!/usr/bin/env python3
"""Audit DOCX/TXT/Markdown for privacy, textual defects, and evidence-language risks."""

from __future__ import annotations

import argparse
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path

from audit_common import PLACEHOLDER_RE, cli_failed
from docx import Document
from docx.oxml.ns import qn

MODES = {"working", "submission", "anonymous", "public"}
SENSITIVE_MODES = {"anonymous", "public"}
NONPORTABLE_FONTS = ("Hiragino", "PingFang", "STHeiti", "Heiti SC", "Songti SC", "Kaiti SC")
PII_PATTERNS = (
    ("中国居民身份证号", re.compile(r"(?<!\d)\d{17}[0-9Xx](?!\d)")),
    ("中国大陆手机号", re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")),
    ("电子邮箱", re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)),
)
KNOWN_DEFECTS = (
    (re.compile(r"(?<!课)堂合作探究"), "堂合作探究", "疑似漏字，应核对是否为“课堂合作探究”"),
    (re.compile(r"(?<!纸)笔测试"), "笔测试", "疑似漏字，应核对是否为“纸笔测试”"),
    (re.compile(r"(?<!小)组讨论"), "组讨论", "疑似漏字，应核对是否为“小组讨论”"),
)
EVIDENCE_PATTERNS = (
    (re.compile(r"从.{0,12}(?:过程)?照片.{0,12}(?:可见|可以看出|看出)"), "照片论据"),
    (re.compile(r"学生(?:们)?普遍(?:认为|反映|表示)"), "群体性学生反馈"),
    (re.compile(r"(?:显著|明显)(?:提高|提升|改善|增强|促进)"), "强效果结论"),
    (re.compile(r"(?:证明了|充分证明|取得了显著|得到广泛推广)"), "证明/推广结论"),
)
NUMBERING_RE = re.compile(r"^\s*(\d{1,2}(?:\.\d{1,2}){1,4})(?=\s|[、.．）)]|[\u4e00-\u9fff])")


def clean(value: str) -> str:
    return re.sub(r"\s+", "", value).strip()


def docx_text(path: Path) -> tuple[list[tuple[str, str]], Document]:
    document = Document(path)
    items: list[tuple[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        if paragraph.text.strip():
            items.append((f"正文段落{index}", paragraph.text.strip()))
    for table_index, table in enumerate(document.tables, 1):
        seen_cells: set[int] = set()
        for row_index, row in enumerate(table.rows, 1):
            for col_index, cell in enumerate(row.cells, 1):
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                value = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if value:
                    items.append((f"表{table_index}行{row_index}列{col_index}", value))
    return items, document


def text_items(path: Path) -> tuple[list[tuple[str, str]], None]:
    value = path.read_text(encoding="utf-8")
    return [(f"第{index}行", line.strip()) for index, line in enumerate(value.splitlines(), 1) if line.strip()], None


def add_privacy_findings(text: str, mode: str, errors: list[str], warnings: list[str]) -> None:
    for label, pattern in PII_PATTERNS:
        count = len(pattern.findall(text))
        if not count:
            continue
        message = f"发现{label}{count}处；请在安全环境核对并制作正确版本（审计输出不显示具体值）"
        if mode in SENSITIVE_MODES:
            errors.append(message)
        else:
            warnings.append(message)


def add_numbering_findings(items: list[tuple[str, str]], warnings: list[str]) -> None:
    by_parent: dict[tuple[int, ...], list[tuple[int, str, str]]] = defaultdict(list)
    seen_full: Counter[tuple[int, ...]] = Counter()
    for location, value in items:
        match = NUMBERING_RE.match(value)
        if not match:
            continue
        number = tuple(int(part) for part in match.group(1).split("."))
        seen_full[number] += 1
        by_parent[number[:-1]].append((number[-1], match.group(1), location))

    for number, count in seen_full.items():
        if count > 1:
            warnings.append(f"标题编号{'.'.join(map(str, number))}重复出现{count}次，请核对目录和正文")
    for entries in by_parent.values():
        encountered: set[int] = set()
        for child, label, location in entries:
            if child > 1 and child - 1 not in encountered:
                warnings.append(f"{location}的编号{label}前缺少同级编号{child - 1}，可能跳号")
            encountered.add(child)


def add_duplicate_findings(items: list[tuple[str, str]], warnings: list[str]) -> None:
    locations: dict[str, list[str]] = defaultdict(list)
    preview: dict[str, str] = {}
    for location, value in items:
        normalized = clean(value)
        if len(normalized) < 50:
            continue
        locations[normalized].append(location)
        preview[normalized] = value
    duplicate_groups = [(key, locs) for key, locs in locations.items() if len(locs) > 1]
    duplicate_groups.sort(key=lambda item: (-len(item[0]), item[1][0]))
    for key, locs in duplicate_groups[:20]:
        snippet = clean(preview[key])[:24]
        warnings.append(f"疑似长段重复：{', '.join(locs[:4])}（“{snippet}…”）")
    if len(duplicate_groups) > 20:
        warnings.append(f"另有{len(duplicate_groups) - 20}组长段重复未逐项列出")


def iter_docx_runs(document: Document):
    for paragraph in document.paragraphs:
        yield from paragraph.runs
    for table in document.tables:
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                for paragraph in cell.paragraphs:
                    yield from paragraph.runs


def add_docx_findings(document: Document, text: str, mode: str, errors: list[str], warnings: list[str]) -> None:
    fonts: Counter[str] = Counter()
    for run in iter_docx_runs(document):
        props = run._r.rPr
        if props is not None:
            fonts_node = props.find(qn("w:rFonts"))
            if fonts_node is not None:
                for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                    value = fonts_node.get(qn(f"w:{key}"))
                    if value:
                        fonts[value] += 1
        if run.font.name:
            fonts[run.font.name] += 1
    risky = {font: count for font, count in fonts.items() if any(token.lower() in font.lower() for token in NONPORTABLE_FONTS)}
    if risky:
        detail = "、".join(f"{font}({count})" for font, count in sorted(risky.items()))
        warnings.append(f"发现可能仅限Mac或跨机不稳定的中文字体：{detail}；按官方模板或替换为可交付字体")

    if re.search(r"(?:过程)?照片.{0,12}(?:可见|可以看出|看出)", text) and not document.inline_shapes:
        warnings.append("正文以照片作论据，但DOCX中未发现内嵌图片；请补充带题注证据或改为证据编号")

    properties = document.core_properties
    metadata = [properties.author, properties.last_modified_by, properties.comments]
    if mode in SENSITIVE_MODES and any(str(value or "").strip() for value in metadata):
        errors.append("匿名/公开版的DOCX文件属性仍含作者、最后修改者或备注，请清理元数据")


def audit(path: Path, mode: str, final: bool, project_title: str | None, forbidden: list[str]) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    suffix = path.suffix.lower()
    if suffix == ".docx":
        items, document = docx_text(path)
    elif suffix in {".txt", ".md"}:
        items, document = text_items(path)
    else:
        return ["仅支持DOCX、TXT和Markdown；旧DOC请先保留原件并转换工作副本"], warnings

    full_text = "\n".join(value for _, value in items)
    add_privacy_findings(full_text, mode, errors, warnings)
    add_numbering_findings(items, warnings)
    add_duplicate_findings(items, warnings)

    for pattern, defect, guidance in KNOWN_DEFECTS:
        count = len(pattern.findall(full_text))
        if count:
            warnings.append(f"发现“{defect}”{count}处：{guidance}")

    for pattern, label in EVIDENCE_PATTERNS:
        count = len(pattern.findall(full_text))
        if count:
            warnings.append(f"发现{label}{count}处；逐项核对真实数据、原始记录或证据编号，并校准结论强度")

    if final:
        placeholders = len(PLACEHOLDER_RE.findall(full_text))
        if placeholders:
            warnings.append(f"最终版发现占位内容{placeholders}处；仅签章/线下填写区可以保留")

    if project_title and project_title not in full_text:
        warnings.append("未找到当前课题规范题目；核对封面、页眉、正文和文件名是否使用统一题目")

    for value in forbidden:
        count = full_text.count(value)
        if count:
            errors.append(f"发现禁止残留词{count}处；请核对旧项目身份、题目或机构信息")

    if document is not None:
        add_docx_findings(document, full_text, mode, errors, warnings)
    return errors, warnings


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("file", type=Path)
    parser.add_argument("--mode", choices=sorted(MODES), default="working")
    parser.add_argument("--final", action="store_true")
    parser.add_argument("--project-title")
    parser.add_argument("--forbid", action="append", default=[], help="禁止出现的旧题目、姓名或机构；可重复使用")
    args = parser.parse_args()
    try:
        errors, warnings = audit(args.file, args.mode, args.final, args.project_title, args.forbid)
    except Exception as exc:
        print(f"读取或审计失败：{exc}", file=sys.stderr)
        return 2
    for item in warnings:
        print(f"警告：{item}")
    for item in errors:
        print(f"错误：{item}")
    if cli_failed(errors, warnings, args.final):
        print(f"内容完整性审计未通过：{len(errors)}个错误，{len(warnings)}个警告")
        return 1
    print(f"内容完整性审计通过：0个错误，{len(warnings)}个警告")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
