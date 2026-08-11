#!/usr/bin/env python3
"""Generate the package-level authenticity, missing-items, and timeline checklist DOCX."""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from build_generic_docx_templates import (
    add_title,
    fill_cell,
    heading,
    set_repeat_header,
    set_table_geometry,
    setup,
    shade,
)
from docx import Document
from docx.shared import Mm

READY = {"ready", "submitted", "archived"}
MISSING_RE = re.compile(r"待填写|待确认|未提供|未知|unknown|tbd", re.I)
LEVEL_ORDER = {"当前阻断": 0, "必须补充": 1, "阶段待办": 2, "建议增强": 3}
DATE_LABELS = {
    "preliminary_research": "前期调查/预研究",
    "application": "申报",
    "approval": "批准立项",
    "opening": "开题",
    "instrument_design": "工具设计",
    "pilot": "预测试",
    "data_collection_start": "正式采集开始",
    "data_collection_end": "正式采集结束",
    "analysis": "数据清理与分析",
    "intervention_start": "干预实施开始",
    "midterm": "中期检查",
    "intervention_end": "干预实施结束",
    "completion": "计划结题",
}
DATE_ORDER = tuple(DATE_LABELS)
PREREQUISITES = {
    "application": "负责人、单位、选题和当年申报要求已确认",
    "approval": "申请获批并取得立项通知/编号",
    "opening": "批准立项在先，开题方案和专家安排已准备",
    "instrument_design": "研究问题、指标和样本边界已确定",
    "pilot": "工具初稿完成，预测试对象和修订记录已准备",
    "data_collection_start": "预测试及修订完成，告知同意和编码规则可用",
    "data_collection_end": "原始记录、样本流转和异常情况持续登记",
    "analysis": "正式采集结束，原始数据完成清理/编码",
    "intervention_start": "诊断分析完成，策略、课例、量规和记录表可用",
    "midterm": "阶段实施、过程记录、问题与调整已形成",
    "intervention_end": "核心策略按计划实施并记录偏离/参与流失",
    "completion": "数据、照片、成果、证明、研究报告和结题材料完整",
}
ROLE_LABELS = {
    "application": "申请书/申报书",
    "anonymous-form": "匿名活页",
    "opening": "开题报告",
    "consent-form": "告知同意/伦理材料",
    "blank-instrument": "研究工具空白版",
    "interview-guide": "访谈提纲",
    "observation-form": "课堂观察表",
    "assessment-tool": "测评工具",
    "codebook": "变量编码本",
    "raw-data": "真实原始数据",
    "completed-record": "已填写原始记录",
    "transcription": "访谈转写/复核稿",
    "data-workbook": "数据工作簿",
    "analysis-report": "诊断/效果分析报告",
    "intervention-plan": "干预实施方案",
    "lesson-plan": "教学设计/活动方案",
    "task-sheet": "学生任务单",
    "rubric": "评价量规",
    "case": "教学案例",
    "casebook": "教学案例集",
    "fidelity-log": "实施保真度记录",
    "meeting-record": "教研/会议记录",
    "progress-report": "中期/阶段报告",
    "photo-register": "照片登记台账",
    "evidence-book": "照片/作品证据册",
    "final-report": "最终研究报告",
    "closing-application": "结题申请/鉴定书",
    "achievement-catalog": "成果目录",
    "expert-appraisal": "专家鉴定/评议",
    "proof": "应用、推广、发表或获奖证明",
    "index": "整套交付索引",
    "attention-items": "材料包注意事项文件",
}

SUBJECT_PROFILES = (
    {
        "label": "心理健康、德育与班主任工作", "keywords": ("心理", "班主任", "德育", "生命教育"),
        "originals": "活动方案、告知同意、匿名量表/访谈、活动记录、转介预案和过程反思",
        "outcomes": "参与和技能表现、匿名趋势、过程反馈；不越界作临床诊断或强因果结论",
        "checks": "工具适用年龄、来源、信效度、计分规则和解释边界",
        "risks": "最小化收集、保密、风险识别和转介；敏感答案/个案不拍摄、不进公开附件",
        "photos": (("环境与准备", "不识别个人的场地、材料和活动准备"), ("集体活动", "经许可的集体活动远景或背影"), ("成果反馈", "匿名成果、反馈工具和教师反思；不拍敏感答案")),
    },
    {
        "label": "信息科技与通用技术", "keywords": ("信息科技", "信息技术", "通用技术", "计算机", "编程", "人工智能"),
        "originals": "源文件/代码、版本记录、操作日志、截图/录屏、测试用例、设备材料清单和作品文件",
        "outcomes": "任务完成、算法/设计过程、调试记录、可用性测试、迭代和迁移",
        "checks": "软件/硬件版本、算法逻辑、数据来源、许可证、模型或平台能力边界",
        "risks": "账号、网络和个人数据安全；开源许可证；AI辅助使用登记；工具加工和用电安全",
        "photos": (("真实操作", "学生在真实设备环境完成任务；避免暴露账号和个人数据"), ("设计制作", "硬件连接、结构制作或作品开发过程"), ("测试迭代", "测试、调试和版本改进过程；截图/日志另作电子证据")),
    },
    {
        "label": "道德与法治、历史、地理", "keywords": ("道德与法治", "思想政治", "政治", "历史", "地理", "思政"),
        "originals": "政策原文、档案/史料、地图、官方统计、访谈/田野记录和议题学习成果",
        "outcomes": "证据辨析、时空定位、材料论证、价值判断过程和地图/图表解读",
        "checks": "文号日期、史实、时间线、行政区划、国家版图、统计口径、来源与时效",
        "risks": "使用权威地图和时政来源；田野访谈知情同意；地方/民族文化尊重来源与社区权益",
        "photos": (("资料研读", "史料、地图或官方数据研读与证据辨析"), ("议题探究", "讨论、实地调查或访谈过程；同步取得来源与同意"), ("成果表达", "地图、时间线、调查报告或议题成果展示")),
    },
    {
        "label": "劳动与综合实践", "keywords": ("劳动", "综合实践", "研究性学习", "社区实践"),
        "originals": "任务方案、材料工具清单、安全培训、分工、过程照片、产品/服务成果和成本记录",
        "outcomes": "规范操作、问题解决、合作、劳动品质、产品质量和反思",
        "checks": "工艺步骤、材料性能、成本/计量、真实社区需求和课程目标",
        "risks": "刀具、热源、食品、校外活动和交通风险；家校/社区同意；合作方与成果署名",
        "photos": (("安全准备", "安全培训、工具材料检查和任务分工"), ("规范操作", "真实劳动/实践过程和协作"), ("产品服务", "成果、使用场景和真实对象反馈")),
    },
    {
        "label": "体育与健康", "keywords": ("体育", "体能", "健康教育", "运动"),
        "originals": "健康筛查/家长同意、运动方案、场地器材检查、原始体测、出勤负荷、视频和异常记录",
        "outcomes": "技术动作量规、体能变化、参与度和自我管理；按个体与分组解释",
        "checks": "动作要领、测量单位、测试流程、负荷与恢复、健康知识边界",
        "risks": "医疗/健康隐私、差异化安排、急救预案、天气场地风险；不作医疗诊断",
        "photos": (("安全热身", "场地器材检查、热身和安全说明"), ("技术练习", "关键动作练习；用视频/量规补足动态证据"), ("测试反馈", "规范测试流程和反馈；不公开个人健康数据")),
    },
    {
        "label": "音乐", "keywords": ("音乐", "声乐", "器乐", "合唱"),
        "originals": "曲谱/资源来源、排练音视频、演唱演奏录音、创编草稿、展演记录和评价量规",
        "outcomes": "节奏、音准、表现、合作、创编过程和修改轨迹",
        "checks": "曲目版本、术语、节拍调式、文化背景和课程标准要求",
        "risks": "曲谱、伴奏、录音、演出和公开传播授权；学生声音与肖像保护",
        "photos": (("排练过程", "分声部/器乐排练和教师指导"), ("创编修改", "创编草稿、合练调整或作品修改"), ("展演评价", "真实展演与量规反馈；音频/视频单独授权登记")),
    },
    {
        "label": "美术", "keywords": ("美术", "书法", "绘画", "设计艺术"),
        "originals": "构思草图、材料试验、创作过程照片、阶段稿—完成稿、作品说明和评价量规",
        "outcomes": "构图、造型、色彩/材料运用、创意过程、修改与表达",
        "checks": "作品/艺术家来源、技法、材料属性、文化背景和课程标准要求",
        "risks": "作品与图片版权；学生作品署名/匿名和展览授权；刀具、颜料、胶黏剂等材料安全",
        "photos": (("构思试验", "草图、材料试验和构思说明"), ("创作过程", "真实制作过程与阶段变化"), ("作品评价", "完成作品、展示和量规反馈；登记作者及授权")),
    },
    {
        "label": "化学", "keywords": ("化学",),
        "originals": "试剂/浓度/用量清单、实验方案、原始现象、测量记录、废液处置和实验照片",
        "outcomes": "现象描述、证据解释、变量控制、定量计算、操作与安全量规",
        "checks": "方程式、条件、配平、物质性质、浓度单位和结论边界",
        "risks": "危化品、明火、通风、PPE、学校制度/MSDS和废弃物处置；不设计不适龄高风险实验",
        "photos": (("实验安全", "试剂器材检查、PPE和安全说明"), ("原始观察", "变量控制、真实现象和测量记录"), ("结果处理", "数据整理、误差分析和废弃物规范处置")),
    },
    {
        "label": "物理", "keywords": ("物理",),
        "originals": "实验方案、器材清单、校准记录、原始测量表、实验照片和误差分析",
        "outcomes": "变量控制、测量质量、数据图表、证据推理和实验操作量规",
        "checks": "公式、单位、有效数字、实验条件、误差边界和模型适用范围",
        "risks": "用电、热、光、机械等风险评估；器材检查、教师监管和异常记录",
        "photos": (("器材安全", "器材准备、校准和风险检查"), ("原始测量", "变量控制、操作与实时读数记录"), ("证据推理", "数据图表、误差分析和模型边界讨论")),
    },
    {
        "label": "生物与小学科学", "keywords": ("生物", "小学科学", "科学"),
        "originals": "观察日志、实验/探究记录、样本或环境照片、测量数据和分类/调查表",
        "outcomes": "观察质量、证据推理、长期变化、探究设计和科学表达",
        "checks": "物种/结构名称、变量、样本条件、健康与生态结论边界",
        "risks": "人体/健康数据隐私，动植物与野外活动伦理，生物材料和器具安全",
        "photos": (("安全准备", "观察/实验材料、场地和安全说明"), ("持续观察", "真实样本、环境、测量与观察记录"), ("探究成果", "分类、数据图表、模型或科学表达成果")),
    },
    {
        "label": "英语", "keywords": ("英语", "外语"),
        "originals": "听说录音/视频、阅读作答、写作初稿—修改稿、词汇/语法诊断和任务量规",
        "outcomes": "听说读看写分项、流利度/准确度/得体性、任务完成度和修改轨迹",
        "checks": "语言表达、语音语调、语法语用、语篇来源、教材版本和课标级别",
        "risks": "音视频与未成年人声音授权；语料、歌曲、绘本、影视和试题版权；公开材料去标识",
        "photos": (("语言任务", "真实交流、朗读或展示；声音证据需另行授权"), ("阅读写作", "文本研读、写作修改和同伴反馈"), ("表现评价", "任务成果与量规反馈；保留音视频/作品原件")),
    },
    {
        "label": "数学", "keywords": ("数学",),
        "originals": "诊断卷、解题草稿、错题原件、操作/建模记录、前后测和评分细则",
        "outcomes": "思路步骤、表征转换、推理质量、错误类型和迁移任务，不只比较总分",
        "checks": "公式、符号、单位、定义、条件、答案等价性和统计口径",
        "risks": "不虚构解题过程；公开学生作答去姓名；数字工具保留版本和导出记录",
        "photos": (("操作探究", "学具操作、建模或问题表征过程"), ("思路交流", "真实板演、讨论和多种解法比较"), ("错因改进", "订正、反馈和迁移任务；原始草稿/测评另行归档")),
    },
    {
        "label": "语文", "keywords": ("语文", "汉语文"),
        "originals": "阅读批注、朗读/口语录音、习作初稿—修改稿、课堂实录、阅读单和教师反馈",
        "outcomes": "阅读理解分项、文本证据使用、表达质量、修改轨迹和口语表现量规",
        "checks": "篇目版本、作者与出处、字词句、文体知识和课程标准行为动词",
        "risks": "诗文、教材截图和影视资源合理使用；学生作文、声音和生活叙事去标识并授权",
        "photos": (("阅读研讨", "批注、讨论和文本证据使用过程"), ("表达修改", "朗读/口语活动或习作修改过程"), ("作品反馈", "初稿—修改稿、展示和量规反馈；声音/作文另行授权")),
    },
)

GENERIC_SUBJECT_PROFILE = {
    "label": "通用学科底座（主学科未识别）",
    "originals": "课程标准/教材定位、真实基线、研究工具、原始记录、实施材料、学生成果和评价反思",
    "outcomes": "与研究问题相称的知识、技能、过程和作品证据，不能只用满意度或活动照片",
    "checks": "课程标准版本、教材范围、概念/术语、评价口径和学段适宜性",
    "risks": "学生隐私、作品/资源版权、活动安全和公开授权；由本学科教师补充专项规则",
    "photos": (("学科任务", "本学科关键学习任务的真实过程"), ("成果形成", "学生成果形成和修改过程"), ("评价反馈", "量规评价、反馈和反思过程")),
}


def profile_for_subject(subject: object) -> dict:
    subject = str(subject or "")
    for profile in SUBJECT_PROFILES:
        if any(keyword in subject for keyword in profile["keywords"]):
            return profile
    return GENERIC_SUBJECT_PROFILE


def subject_profile(data: dict) -> dict:
    project = data.get("project", {}) if isinstance(data.get("project"), dict) else {}
    return profile_for_subject(project.get("subject"))


@dataclass
class Issue:
    level: str
    category: str
    item: str
    why: str
    action: str
    target: str


def parse_day(value: object) -> date | None:
    if not value:
        return None
    try:
        return date.fromisoformat(str(value))
    except ValueError:
        return None


def text_missing(value: object) -> bool:
    return value is None or not str(value).strip() or bool(MISSING_RE.search(str(value)))


def display(value: object, fallback: str = "未确认") -> str:
    return fallback if text_missing(value) else str(value)


def material_name(item: dict) -> str:
    return str(item.get("name") or ROLE_LABELS.get(str(item.get("material_role")), item.get("id", "未编号材料")))


def issue_level_for_material(item: dict, delivery_state: str, truth_state: str) -> str:
    status = str(item.get("status", "planned"))
    if item.get("required_for_submission") is True:
        return "当前阻断"
    if delivery_state == "closing-ready" or truth_state in {"closing", "completed"}:
        return "当前阻断"
    if status in {"pending-data", "pending-photo", "pending-signature"}:
        return "必须补充"
    return "阶段待办"


def gather_issues(data: dict, as_of: date) -> list[Issue]:
    issues: list[Issue] = []
    generation = data.get("generation_contract", {}) if isinstance(data.get("generation_contract"), dict) else {}
    delivery_state = str(generation.get("delivery_state", "full-lifecycle-scaffold"))
    truth_state = str(generation.get("truth_state", "planning"))
    project = data.get("project", {}) if isinstance(data.get("project"), dict) else {}
    for key, label in (("title", "规范题目"), ("leader", "负责人"), ("school", "学校全称"), ("subject", "学科"), ("stage", "学段")):
        if text_missing(project.get(key)):
            issues.append(Issue("当前阻断", "项目身份", f"{label}尚未确认", "该字段会同步影响所有材料和报送资格", f"由课题负责人核对正式写法并更新项目主清单；完成日期不晚于{as_of.isoformat()}", "全部材料；完成标准：主清单中只有一个正式值"))

    requirements = data.get("submission_requirements", {}) if isinstance(data.get("submission_requirements"), dict) else {}
    req_status = str(requirements.get("status", "pending"))
    if req_status != "verified":
        level = "当前阻断" if delivery_state in {"application-ready", "closing-ready"} else "必须补充"
        issues.append(Issue(level, "官方要求", f"当年通知和官方模板状态为{req_status}", "年度栏目、匿名、限额、签章、命名和截止时间可能变化", "取得管理单位当年通知及全部附件，登记来源、适用年度、核验日期和截止日期；责任人：课题负责人/学校科研管理人员", "申请书、开题、结题及上传目录；完成标准：requirements.status=verified"))

    for source in data.get("sources", []):
        if not isinstance(source, dict) or source.get("verification_status") == "verified":
            continue
        issues.append(Issue("必须补充", "来源核验", f"来源{source.get('id', '未编号')}《{source.get('title', '未命名')}》尚未核验", "政策、课程标准、官方附件或文献必须可追溯", f"从权威渠道取得原件并核对发布日期、文号、适用年度和定位；责任人：{source.get('owner', '课题负责人')}", f"使用位置：{source.get('used_in', []) or '主清单登记位置'}；完成标准：verification_status=verified"))

    materials = [item for item in data.get("materials", []) if isinstance(item, dict)]
    for item in materials:
        if item.get("material_role") == "attention-items" or item.get("status") in READY:
            continue
        status = str(item.get("status", "planned"))
        level = issue_level_for_material(item, delivery_state, truth_state)
        action_by_status = {
            "pending-data": "取得真实原始记录/数据，完成录入、清理、编码和复核后再生成结果",
            "pending-photo": "取得真实照片原件、日期地点、拍摄来源和授权后再插入并生成图题",
            "pending-signature": "由对应负责人、专家或单位在线下完成真实签字/盖章",
            "draft": "核对内容、格式、来源和依赖后完成审计",
            "verified": "完成文件制作、哈希登记和各项QA后标记ready",
            "planned": "按时间轴生成并在真实活动发生后补齐证据",
        }
        issues.append(Issue(level, "材料缺件", f"{item.get('id', '未编号')} {material_name(item)}当前状态为{status}", "材料尚未达到当前或后续阶段的可用门槛", f"{action_by_status.get(status, '核对并完成该材料')}；责任人：{item.get('owner', '课题负责人')}；最迟时间：{item.get('due_date') or '按项目时间轴相应节点前'}", f"材料ID：{item.get('id', '未编号')}；完成标准：文件存在、状态ready、QA及哈希可追溯"))

    for contributor in data.get("contributors", []):
        if isinstance(contributor, dict) and contributor.get("confirmed") is False:
            contributor_label = display(contributor.get("name"), str(contributor.get("id") or "未编号成员"))
            issues.append(Issue("必须补充", "人员与署名", f"成员{contributor_label}的身份/排序/贡献尚未确认", "成员资格和成果署名错误会影响申报与结题", "核对姓名、单位、角色、成员排序、实际贡献和是否为获批成员；由负责人和本人共同确认", f"成员ID：{contributor.get('id', '未编号')}；完成标准：confirmed=true并有贡献登记"))

    timeline = data.get("timeline", {}) if isinstance(data.get("timeline"), dict) else {}
    completion = parse_day(timeline.get("completion"))
    governance = data.get("governance", {}) if isinstance(data.get("governance"), dict) else {}
    if completion and as_of > completion and governance.get("project_status") not in {"completed", "extended", "terminated"}:
        issues.append(Issue("当前阻断", "时间状态", f"当前日期{as_of.isoformat()}已晚于计划结题日{completion.isoformat()}", "继续沿用旧计划会造成申报、实施和结题时间矛盾", "核实已结题、获批延期、终止或逾期未办状态；如延期，取得批准依据和新期限，不得倒改历史日期", "治理状态和全部结题材料；完成标准：project_status及批准记录一致"))

    previous_key, previous_day = None, None
    for key in DATE_ORDER:
        current_day = parse_day(timeline.get(key))
        if current_day and previous_day and current_day < previous_day:
            issues.append(Issue("当前阻断", "时间倒序", f"{DATE_LABELS[key]}日期{current_day}早于{DATE_LABELS[previous_key]}日期{previous_day}", "研究活动前置关系不成立", "核对真实记录；如属于前期基础则明确标注，涉及正式变更时补批准文件，不得为通过审计而倒改日期", f"timeline.{previous_key} → timeline.{key}；完成标准：真实时间和说明一致"))
        if current_day:
            previous_key, previous_day = key, current_day

    samples = [item for item in data.get("samples", []) if isinstance(item, dict)]
    collection_end = parse_day(timeline.get("data_collection_end"))
    for sample in samples:
        if sample.get("actual_n") is None:
            level = "当前阻断" if truth_state in {"closing", "completed"} or (collection_end and collection_end < as_of) else "阶段待办"
            issues.append(Issue(level, "真实样本", f"{sample.get('name', '样本')}只有计划样本量，尚无实际有效样本量", "计划人数不能作为实际N值写入分析报告", "完成回收、无效判定、去重和样本流转记录后填写actual_n及data_source", "数据工作簿和分析报告；完成标准：实际N值可从原始数据复核"))

    for instrument in data.get("instruments", []):
        if not isinstance(instrument, dict):
            continue
        status = str(instrument.get("status", "draft"))
        if status in {"collected", "analyzed", "completed"} and not instrument.get("raw_evidence"):
            issues.append(Issue("当前阻断", "原始记录", f"工具{instrument.get('id', '未编号')}已标记{status}但未登记raw_evidence", "无法从分析结论回溯到真实问卷、访谈、观察或测评原件", "补登记原始文件、记录ID、页码/数据区域、复核人和隐私位置", f"工具：{instrument.get('name', instrument.get('id'))}；完成标准：raw_evidence可访问并与分析对应"))

    for case in data.get("cases", []):
        if not isinstance(case, dict) or case.get("status") not in {"implemented", "validated"}:
            continue
        missing = []
        implementation = case.get("implementation") if isinstance(case.get("implementation"), dict) else {}
        for key, label in (("date", "日期"), ("school", "学校"), ("class", "班级"), ("teacher", "教师"), ("student_n", "人数"), ("actual_periods", "实际课时")):
            if text_missing(implementation.get(key)):
                missing.append(label)
        if missing or not case.get("evidence_ids"):
            issues.append(Issue("当前阻断", "案例真实性", f"案例{case.get('id', '未编号')}标记为{case.get('status')}但缺少{','.join(missing) or '证据ID'}", "实施性案例必须有可核验的真实课堂元数据和证据", "补实施记录、教案版本、偏离事项、照片/作品/观察/评价ID；无法补齐时降级为designed或piloted", f"案例：{case.get('title', case.get('id'))}；完成标准：实施字段和证据ID完整"))

    for commitment in data.get("commitments", []):
        if not isinstance(commitment, dict) or commitment.get("status") in {"fulfilled", "changed-approved"}:
            continue
        due = parse_day(commitment.get("due_date"))
        level = "当前阻断" if truth_state in {"closing", "completed"} or (due and due < as_of) else "阶段待办"
        issues.append(Issue(level, "成果承诺", f"承诺成果{commitment.get('id', '未编号')}《{commitment.get('name', '未命名')}》状态为{commitment.get('status', 'planned')}", "结题必须逐项兑现申请书承诺或提供获批变更", "形成真实成果并登记材料/证明；确需变更时办理批准，不能静默改名或删除", f"承诺ID：{commitment.get('id', '未编号')}；完成标准：fulfilled或changed-approved"))

    deduped: list[Issue] = []
    seen: set[tuple[str, str, str]] = set()
    for issue in issues:
        key = (issue.level, issue.category, issue.item)
        if key not in seen:
            seen.add(key)
            deduped.append(issue)
    return sorted(deduped, key=lambda item: (LEVEL_ORDER[item.level], item.category, item.item))


def gather_photo_rows(data: dict, as_of: date) -> list[list[str]]:
    materials = {str(item.get("id")): material_name(item) for item in data.get("materials", []) if isinstance(item, dict)}
    photos = [item for item in data.get("evidence", []) if isinstance(item, dict) and item.get("type") == "photo"]
    rows: list[list[str]] = []
    for index, item in enumerate(photos, 1):
        photo_id = str(item.get("id") or f"PHO-{as_of.year}-{index:03d}")
        missing = [label for key, label in (("collected_date", "日期"), ("location", "地点"), ("photographer_source", "拍摄/来源"), ("source_file", "原件"), ("original_sha256", "哈希"), ("caption", "图题"), ("alt_text", "替代文本")) if text_missing(item.get(key))]
        target_ids = [str(value) for value in item.get("material_ids", [])]
        targets = "、".join(f"{value} {materials.get(value, '')}".strip() for value in target_ids) or "照片证据册、对应案例/报告"
        activity = str(item.get("activity") or "按研究节点拍摄真实活动")
        consent = f"同意：{item.get('consent_status', '未登记')}；人物处理：{item.get('face_handling', '未登记')}"
        status = str(item.get("status", "planned"))
        if missing:
            status += "；缺" + "、".join(missing)
        rows.append([photo_id, "按活动日期归属阶段", activity, "真实日期、地点、拍摄来源、对应记录/案例ID、原件文件和SHA-256", consent, targets, status])
    if rows:
        return rows
    profile = subject_profile(data)
    suggestions = [
        ("研究准备", "集体研讨、工具设计或预测试现场", "开题/过程记录/中期"),
        ("正式采集", "问卷、访谈、课堂观察或测评的真实过程（避免拍到敏感答案）", "工具记录/中期"),
    ]
    suggestions.extend((f"学科专项：{stage}", activity, "教案/案例/学科证据册/研究报告") for stage, activity in profile["photos"])
    suggestions.extend((
        ("评价改进", "量规评价、反馈讨论、修改前后成果或教学反思", "案例/中期/研究报告"),
        ("应用推广", "校内交流、成果使用、培训或推广活动", "结题证明/证据册"),
    ))
    for index, (stage, activity, target) in enumerate(suggestions, 1):
        rows.append([f"PHO-{as_of.year}-{index:03d}", stage, f"建议拍摄：{activity}；仅在真实发生时取得", "日期、地点、活动名称、拍摄者、对应记录ID、原件和哈希", "优先背影/远景/作品特写；公开使用前核对授权并打码", target, "建议任务；尚未登记真实照片"])
    return rows


def role_status(data: dict, roles: set[str]) -> str:
    matching = [item for item in data.get("materials", []) if isinstance(item, dict) and item.get("material_role") in roles]
    if not matching:
        return "未登记"
    if any(item.get("status") in READY for item in matching):
        return "已有就绪文件"
    return "、".join(sorted({str(item.get("status", "planned")) for item in matching}))


def gather_original_rows(data: dict) -> list[list[str]]:
    definitions = (
        ("官方与审批", {"official-form", "application", "opening", "closing-application"}, "通知、指南、官方表格、立项/开题/中期/变更/延期/结题批复", "保留完整原件、文号、发布日期、适用年度、签章和来源定位", "01_官方与审批", "申报、立项、开题、变更和结题节点"),
        ("伦理与授权", {"consent-form"}, "教师/学生/家长告知同意、照片和公开使用授权", "按学校制度取得；公开版不附敏感授权原件，只登记状态", "02_伦理与授权_受控", "正式采集、拍摄和公开前"),
        ("工具与记录", {"blank-instrument", "completed-record", "interview-guide", "observation-form", "transcription"}, "空白工具、已填问卷、访谈录音/转写、观察原表和复核记录", "空白版与填写版分开；每份有记录ID、日期、对象、页码/录音定位和复核人", "03_工具与原始记录_受控", "预测试、正式采集和分析前"),
        ("课堂与作品", {"lesson-plan", "task-sheet", "rubric", "case"}, "实施教案、任务单、学生作品、量规评分、观察和反思", "日期、班级、教师、人数、课时、版本和证据ID一致；作品去姓名", "04_课堂实践与学生作品", "每次课例实施当日或随后及时整理"),
        ("数据与分析", {"codebook", "raw-data", "data-workbook", "analysis-report"}, "变量编码、原始数据、清理记录、公式分析和结果表", "原始数据只增补不覆盖；分析可追溯到工作表区域、N值、分母和截止日期", "05_数据与分析_受控", "采集后、分析报告前"),
        ("过程管理", {"meeting-record", "fidelity-log", "progress-report", "change-form"}, "会议、集体备课、实施保真度、中期、问题调整和变更记录", "记录真实日期、人员、议题、决定、偏离原因和附件证据", "06_过程管理", "活动发生后及时形成"),
        ("成果与推广", {"final-report", "casebook", "achievement-catalog", "proof", "expert-appraisal"}, "研究报告、案例集、论文、获奖、发表、应用反馈、推广和专家鉴定", "成果作者与贡献一致；证明含日期、单位、对象、使用范围和可核验来源", "07_成果与推广", "成果形成、应用和结题阶段"),
        ("签章", {"application", "opening", "closing-application", "expert-appraisal"}, "负责人签名、成员确认、专家意见、学校及管理部门盖章", "必须由真人/单位完成；不复制其他项目签章，不预填专家意见", "08_签章与报送原件_受控", "按官方流程和截止日期"),
    )
    return [[category, original, standard, folder, timing, role_status(data, roles)] for category, roles, original, standard, folder, timing in definitions]


def gather_subject_rows(data: dict) -> list[list[str]]:
    project = data.get("project", {}) if isinstance(data.get("project"), dict) else {}
    subjects = [(display(project.get("subject"), "主学科未确认"), "主学科")]
    related = project.get("related_subjects", [])
    if isinstance(related, list):
        subjects.extend((display(value, "未确认关联学科"), "关联学科") for value in related if str(value).strip())
    completion = "在工具、课例、证据册、数据表和研究报告中登记对应材料ID；原件、版本、日期、对象、评分/分析口径可追溯"
    rows: list[list[str]] = []
    seen: set[str] = set()
    for subject, role in subjects:
        if subject in seen:
            continue
        seen.add(subject)
        profile = profile_for_subject(subject)
        standard = completion
        if profile is GENERIC_SUBJECT_PROFILE:
            standard += "；负责人须补充本学科课程标准、典型学习成果、评价量规和专项风险后再进入实施阶段"
        rows.append([f"{role}：{subject}｜{profile['label']}", profile["originals"], profile["outcomes"], profile["checks"], profile["risks"], standard])
    return rows


def gather_timeline_rows(data: dict, as_of: date) -> list[list[str]]:
    timeline = data.get("timeline", {}) if isinstance(data.get("timeline"), dict) else {}
    rows: list[list[str]] = []
    previous_day: date | None = None
    for key in DATE_ORDER:
        planned = parse_day(timeline.get(key))
        if not planned:
            continue
        if previous_day and planned < previous_day:
            result = "冲突：早于前置节点"
            advice = "核对真实日期；前期活动单独标注，正式变更补批准依据"
        elif planned < as_of:
            result = "计划节点已过，核对真实证据"
            advice = "登记实际日期、原始记录/批复/材料ID；未完成则如实说明延期或调整"
        elif planned == as_of:
            result = "本日节点"
            advice = "同步保留当日原始记录、照片、参与人和版本"
        else:
            result = "计划节点未到"
            advice = "提前准备前置材料；不得预写实际结果"
        rows.append([DATE_LABELS[key], planned.isoformat(), "以真实材料/证据ID登记", PREREQUISITES.get(key, "按项目主清单前置关系"), result, advice, f"不晚于{planned.isoformat()}；课题负责人/对应成员"])
        previous_day = planned
    if not rows:
        rows.append(["完整研究时间轴", "未登记", "无", "申报、立项、开题、工具、采集、分析、实施、中期、结题", "当前阻断", "先建立真实可执行时间轴，再批量生成材料", f"尽快；课题负责人（检查日期{as_of.isoformat()}）"])
    return rows


def gather_suggestions(data: dict) -> list[list[str]]:
    suggestions: list[list[str]] = []
    role_sets = {str(item.get("material_role")) for item in data.get("materials", []) if isinstance(item, dict)}
    def add(priority: str, item: str, reason: str, method: str, stage: str) -> None:
        suggestions.append([priority, item, reason, method, stage])
    if "consent-form" not in role_sets:
        add("A", "告知同意与照片授权材料", "涉及学生数据、作品和可识别人物，缺授权会限制报送和公开", "按学校制度制作教师/学生/家长版本并登记适用范围", "正式采集和拍摄前；实施必需")
    if "codebook" not in role_sets:
        add("A", "变量编码本和题项—指标矩阵", "保证问卷、XLSX、统计表和Word结论使用同一口径", "登记变量名、标签、类型、取值、缺失值、题号、维度和分析方法", "工具定稿前；数据分析必需")
    if "fidelity-log" not in role_sets:
        add("A", "实施保真度与偏离记录", "策略没有按计划实施时，效果结论容易失真", "逐次记录核心成分、课时、覆盖对象、完成度、偏离原因和证据ID", "干预实施阶段；结题解释重要")
    if "progress-report" not in role_sets:
        add("B", "中期报告和阶段问题调整表", "形成申请—开题—中期—结题的递进链，保留真实调整", "列完成情况、阶段数据、问题、变更、后续计划和证明", "中期节点；视管理要求")
    if not any(item.get("type") == "student-work" for item in data.get("evidence", []) if isinstance(item, dict)):
        add("A", "学生作品样本及评价量规", "仅有活动照片不能证明学生层面变化", "分层抽取作品，去标识，保留评分量规、评分人、复核和前后版本", "实施与效果评价；效果结论关键")
    if not any(item.get("type") == "dissemination" for item in data.get("evidence", []) if isinstance(item, dict)):
        add("B", "成果应用、使用反馈和推广证明", "增强成果价值与可推广性证据", "取得使用单位/教师反馈、应用日期、对象、范围、版本和真实附件", "成果应用后；通常非申报阻断")
    add("B", "预测试与工具修订记录", "证明题项、观察指标和量规经过可用性检查", "保留预测试对象、问题、修改前后版本、修订原因和确认人", "正式采集前；方法质量增强")
    add("B", "反例、无效案例和研究局限记录", "避免只呈现成功案例和过度归因，提高研究可信度", "登记未达预期活动、异常数据、参与流失、替代解释和改进", "实施至结题；非硬性附件但强烈建议")
    return suggestions


def add_dynamic_table(doc: Document, headers: tuple[str, ...], widths: list[int], rows: list[list[str]], center_cols: set[int] | None = None) -> None:
    center_cols = center_cols or set()
    table = doc.add_table(rows=max(2, len(rows) + 1), cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for col, value in enumerate(headers):
        fill_cell(table.rows[0].cells[col], value, True, True, 9)
        shade(table.rows[0].cells[col])
    effective_rows = rows or [["无"] + ["—"] * (len(headers) - 1)]
    for r, row_values in enumerate(effective_rows, 1):
        for c, value in enumerate(row_values):
            fill_cell(table.rows[r].cells[c], str(value), center=c in center_cols, size=9)


def build_document(data: dict, output: Path, as_of: date) -> Counter:
    project = data.get("project", {}) if isinstance(data.get("project"), dict) else {}
    governance = data.get("governance", {}) if isinstance(data.get("governance"), dict) else {}
    generation = data.get("generation_contract", {}) if isinstance(data.get("generation_contract"), dict) else {}
    requirements = data.get("submission_requirements", {}) if isinstance(data.get("submission_requirements"), dict) else {}
    issues = gather_issues(data, as_of)
    counts = Counter(issue.level for issue in issues)

    doc = Document(); setup(doc, "课题材料包注意事项与真实性待办清单")
    section = doc.sections[0]
    section.page_width = Mm(297); section.page_height = Mm(210)
    section.top_margin = Mm(18); section.bottom_margin = Mm(18); section.left_margin = Mm(18); section.right_margin = Mm(18)
    add_title(doc, f"{display(project.get('title'), '课题')}材料包注意事项", f"真实性、缺件、时间逻辑与增强建议｜检查日期：{as_of.isoformat()}")

    callout = doc.add_table(rows=1, cols=1); set_table_geometry(callout, [14520])
    callout_text = "本文件供课题负责人内部核对，不代替官方申报/结题附件。照片、数据、原始记录、证明、专家意见、签字和盖章必须来自真实项目；网络图、AI图和其他项目材料不得冒充。"
    fill_cell(callout.rows[0].cells[0], callout_text, size=10.5); shade(callout.rows[0].cells[0], "FFF2CC")

    heading(doc, "一、材料包状态快照")
    overview_rows = [
        ["课题/负责人", f"{display(project.get('title'))} / {display(project.get('leader'))}", "单位/项目状态", f"{display(project.get('school'))} / {display(governance.get('project_status'), '未登记')}"],
        ["成套范围/真实性", f"{generation.get('package_scope', '未登记')} / {generation.get('truth_state', '未登记')}", "目标状态", str(generation.get('delivery_state', '未登记'))],
        ["事项统计", f"当前阻断{counts['当前阻断']}；必须补充{counts['必须补充']}；阶段待办{counts['阶段待办']}", "数据/要求快照", f"数据截止：{governance.get('data_cutoff') or '未登记'}；官方要求：{requirements.get('status', '未登记')}"],
    ]
    add_dynamic_table(doc, ("项目", "内容", "控制项", "状态"), [1900, 5360, 1900, 5360], overview_rows, {0, 2})

    heading(doc, "二、缺失、矛盾与待人工完成事项")
    issue_rows = []
    for index, issue in enumerate(issues, 1):
        issue_rows.append([f"ATT-{index:03d}", issue.level, issue.category, issue.item, issue.why, issue.action, issue.target])
    if not issue_rows:
        issue_rows = [["ATT-000", "已核验", "总体", "当前未发现阻断或必补缺件", "仍需按真实活动持续更新", "每次新增数据、照片、证明或变更后重建本文件", "以最新主清单和一键预检为准"]]
    add_dynamic_table(doc, ("ID", "级别", "类别", "具体事项", "为什么需要", "建议行动/责任/最迟时间", "对应材料与完成标准"), [650, 900, 1300, 3000, 2500, 3600, 2570], issue_rows, {0, 1})

    heading(doc, "三、真实照片拍摄、提供与插入清单")
    add_dynamic_table(doc, ("照片ID", "阶段", "真实活动/建议画面", "必须同步取得的信息", "授权/人物处理", "目标材料与插入位置", "最迟时间/状态"), [850, 1200, 2600, 2800, 2200, 2700, 2170], gather_photo_rows(data, as_of), {0, 1})

    heading(doc, "四、其他真实原始材料清单")
    add_dynamic_table(doc, ("类别", "真实原件", "最低标准", "命名/保存目录", "对应时间与用途", "当前状态"), [1200, 2200, 3400, 2400, 2700, 2620], gather_original_rows(data), {0})

    heading(doc, "五、学科专项真实性与安全核验")
    add_dynamic_table(doc, ("学科/领域", "应取得的真实材料", "学习或效果证据", "关键学科事实", "安全/伦理/版权", "对应材料与完成标准"), [1300, 2800, 2500, 2500, 2900, 2520], gather_subject_rows(data), {0})

    heading(doc, "六、时间逻辑检查")
    add_dynamic_table(doc, ("节点", "计划日期", "实际日期/证据", "前置条件", "检查结果", "处理建议", "最迟完成时间/责任人"), [1500, 1400, 1800, 2400, 1500, 3200, 2720], gather_timeline_rows(data, as_of), {0, 1, 4})

    heading(doc, "七、建议增加的增强材料")
    suggestion_rows = gather_suggestions(data)
    add_dynamic_table(doc, ("优先级", "建议材料", "增加理由", "建议做法/最低标准", "适用阶段/是否阻断"), [900, 2200, 4400, 4000, 3020], suggestion_rows, {0})
    counts["建议增强"] = len(suggestion_rows)

    heading(doc, "八、教师最终操作顺序")
    top_actions = [issue for issue in issues if issue.level in {"当前阻断", "必须补充"}][:8]
    action_rows: list[list[str]] = []
    for index, issue in enumerate(top_actions, 1):
        action_rows.append([str(index), issue.item, issue.target, issue.action, "完成后更新主清单并重建本文件"])
    if not action_rows:
        action_rows = [
            ["1", "核对当年官方要求和提交系统", "通知、附件、截止时间、命名和签章", "由负责人/学校科研管理人员再次核验", "核验完成后再上传"],
            ["2", "新增真实活动后同步归档", "原始记录、数据、照片、作品、评价和反思", "按证据ID登记并更新材料", "每个活动完成后"],
            ["3", "运行一键预检并人工逐页检查", "DOCX、XLSX、PDF、照片、隐私和签章", "处理全部错误和警告", "正式交付前"],
        ]
    add_dynamic_table(doc, ("顺序", "行动", "需要核对/取得的真实输入", "完成标准或操作", "更新要求"), [800, 3000, 4200, 3800, 2720], action_rows, {0})

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return counts


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("manifest", type=Path)
    parser.add_argument("--root", type=Path, help="材料包根目录；默认使用manifest所在目录")
    parser.add_argument("--out", type=Path, help="输出DOCX；默认根目录/00_课题材料包注意事项_真实性与待办清单.docx")
    parser.add_argument("--as-of", type=str, help="检查日期YYYY-MM-DD；默认使用governance.current_date或今天")
    args = parser.parse_args()
    data = json.loads(args.manifest.read_text(encoding="utf-8"))
    governance = data.get("governance", {}) if isinstance(data.get("governance"), dict) else {}
    as_of = parse_day(args.as_of or governance.get("current_date")) or date.today()
    root = (args.root or args.manifest.parent).resolve()
    output = args.out.resolve() if args.out else root / "00_课题材料包注意事项_真实性与待办清单.docx"
    counts = build_document(data, output, as_of)
    print(output)
    print("；".join(f"{key}{counts[key]}" for key in ("当前阻断", "必须补充", "阶段待办", "建议增强")))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
