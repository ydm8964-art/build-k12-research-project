#!/usr/bin/env python3
"""Extract effective typography, paragraph, page, and table-format evidence from DOCX files."""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def rounded(value: float | None, digits: int = 2) -> float | None:
    return None if value is None else round(value, digits)


def length_pt(value) -> float | None:
    return rounded(value.pt) if value is not None else None


def length_cm(value) -> float | None:
    return rounded(value.cm) if value is not None else None


def style_chain(style) -> list:
    result = []
    seen: set[str] = set()
    while style is not None and style.style_id not in seen:
        result.append(style)
        seen.add(style.style_id)
        style = style.base_style
    return result


def rfonts_name(element, key: str) -> str | None:
    if element is None:
        return None
    rpr = getattr(element, "rPr", None)
    if rpr is None:
        return None
    fonts = rpr.find(qn("w:rFonts"))
    return fonts.get(qn(f"w:{key}")) if fonts is not None else None


def first_not_none(values: list[Any]) -> Any:
    return next((value for value in values if value is not None), None)


def alignment_xml(paragraph_or_style) -> str | None:
    element = paragraph_or_style._element
    ppr = element.pPr if hasattr(element, "pPr") else None
    if ppr is None:
        return None
    jc = ppr.find(qn("w:jc"))
    return jc.get(qn("w:val")) if jc is not None else None


def effective_run(run, paragraph) -> dict[str, Any]:
    styles = style_chain(paragraph.style)
    run_style = style_chain(run.style) if run.style is not None else []
    candidates = [run.font, *(item.font for item in run_style), *(item.font for item in styles)]
    east_asia = first_not_none(
        [rfonts_name(run._element, "eastAsia")]
        + [rfonts_name(item._element, "eastAsia") for item in run_style]
        + [rfonts_name(item._element, "eastAsia") for item in styles]
    )
    ascii_name = first_not_none(
        [rfonts_name(run._element, "ascii"), run.font.name]
        + [rfonts_name(item._element, "ascii") or item.font.name for item in run_style]
        + [rfonts_name(item._element, "ascii") or item.font.name for item in styles]
    )
    size = first_not_none([length_pt(item.size) for item in candidates])
    bold = first_not_none([item.bold for item in candidates])
    italic = first_not_none([item.italic for item in candidates])
    return {
        "font_east_asia": east_asia,
        "font_ascii": ascii_name,
        "size_pt": size,
        "bold": bold,
        "italic": italic,
    }


def effective_paragraph(paragraph) -> dict[str, Any]:
    styles = style_chain(paragraph.style)
    formats = [paragraph.paragraph_format, *(item.paragraph_format for item in styles)]
    alignment = first_not_none([alignment_xml(paragraph), *(alignment_xml(item) for item in styles)])
    line_spacing = first_not_none([item.line_spacing for item in formats])
    if hasattr(line_spacing, "pt"):
        line_spacing_value: float | str | None = f"{rounded(line_spacing.pt)}pt"
    elif line_spacing is None:
        line_spacing_value = None
    else:
        line_spacing_value = rounded(float(line_spacing))
    return {
        "style": paragraph.style.name,
        "alignment": alignment,
        "first_line_indent_pt": length_pt(first_not_none([item.first_line_indent for item in formats])),
        "left_indent_pt": length_pt(first_not_none([item.left_indent for item in formats])),
        "right_indent_pt": length_pt(first_not_none([item.right_indent for item in formats])),
        "space_before_pt": length_pt(first_not_none([item.space_before for item in formats])),
        "space_after_pt": length_pt(first_not_none([item.space_after for item in formats])),
        "line_spacing": line_spacing_value,
        "keep_with_next": first_not_none([item.keep_with_next for item in formats]),
        "page_break_before": first_not_none([item.page_break_before for item in formats]),
    }


def paragraph_role(text: str, record: dict[str, Any], ordinal: int) -> str:
    style = record["paragraph"]["style"].lower()
    if "title" in style or (ordinal <= 4 and record["paragraph"]["alignment"] in {"center", "centerContinuous"}):
        return "title"
    if "heading 1" in style or re.match(r"^[一二三四五六七八九十]+[、．.]", text):
        return "heading1"
    if "heading 2" in style or re.match(r"^[（(][一二三四五六七八九十]+[）)]", text):
        return "heading2"
    if "heading 3" in style or re.match(r"^\d+[.、．]", text):
        return "heading3"
    if re.match(r"^[（(]\d+[）)]", text):
        return "heading4"
    return "body"


def record_paragraph(paragraph, ordinal: int, location: str) -> dict[str, Any] | None:
    text = paragraph.text.strip()
    if not text:
        return None
    visible_runs = [run for run in paragraph.runs if run.text.strip()]
    runs = [effective_run(run, paragraph) | {"chars": len(run.text.strip())} for run in visible_runs]
    if runs:
        dominant = Counter(
            (
                item["font_east_asia"], item["font_ascii"], item["size_pt"],
                item["bold"], item["italic"],
            )
            for item in runs
            for _ in range(max(item["chars"], 1))
        ).most_common(1)[0][0]
        run_format = dict(zip(("font_east_asia", "font_ascii", "size_pt", "bold", "italic"), dominant))
    else:
        run_format = {"font_east_asia": None, "font_ascii": None, "size_pt": None, "bold": None, "italic": None}
    record = {
        "location": location,
        "text": text[:160],
        "paragraph": effective_paragraph(paragraph),
        "run": run_format,
    }
    record["role"] = paragraph_role(text, record, ordinal)
    return record


def compact_counts(records: list[dict[str, Any]], key) -> list[dict[str, Any]]:
    counter = Counter(key(item) for item in records)
    return [{"value": value, "count": count} for value, count in counter.most_common(12)]


def extract(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs: list[dict[str, Any]] = []
    ordinal = 0
    for paragraph in document.paragraphs:
        ordinal += 1
        record = record_paragraph(paragraph, ordinal, "body")
        if record:
            paragraphs.append(record)
    table_records: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, 1):
        cell_paragraphs: list[dict[str, Any]] = []
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                for paragraph in cell.paragraphs:
                    ordinal += 1
                    record = record_paragraph(paragraph, ordinal, f"table-{table_index}-r{row_index}c{cell_index}")
                    if record:
                        cell_paragraphs.append(record)
        grid = table._tbl.tblGrid
        widths = [int(node.get(qn("w:w"))) for node in grid.findall(qn("w:gridCol")) if node.get(qn("w:w"))]
        table_records.append({
            "index": table_index,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "grid_widths_dxa": widths,
            "paragraph_count": len(cell_paragraphs),
            "font_size_counts": compact_counts(cell_paragraphs, lambda item: item["run"]["size_pt"]),
            "font_counts": compact_counts(cell_paragraphs, lambda item: item["run"]["font_east_asia"] or item["run"]["font_ascii"]),
            "alignment_counts": compact_counts(cell_paragraphs, lambda item: item["paragraph"]["alignment"]),
            "sample": cell_paragraphs[:12],
        })
    sections = []
    for section in document.sections:
        sections.append({
            "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
            "page_width_cm": length_cm(section.page_width),
            "page_height_cm": length_cm(section.page_height),
            "top_margin_cm": length_cm(section.top_margin),
            "bottom_margin_cm": length_cm(section.bottom_margin),
            "left_margin_cm": length_cm(section.left_margin),
            "right_margin_cm": length_cm(section.right_margin),
            "header_distance_cm": length_cm(section.header_distance),
            "footer_distance_cm": length_cm(section.footer_distance),
        })
    summaries = {}
    for role in ("title", "heading1", "heading2", "heading3", "heading4", "body"):
        selected = [item for item in paragraphs if item["role"] == role]
        summaries[role] = {
            "count": len(selected),
            "fonts": compact_counts(selected, lambda item: item["run"]["font_east_asia"] or item["run"]["font_ascii"]),
            "sizes": compact_counts(selected, lambda item: item["run"]["size_pt"]),
            "bold": compact_counts(selected, lambda item: item["run"]["bold"]),
            "alignment": compact_counts(selected, lambda item: item["paragraph"]["alignment"]),
            "first_line_indent_pt": compact_counts(selected, lambda item: item["paragraph"]["first_line_indent_pt"]),
            "line_spacing": compact_counts(selected, lambda item: item["paragraph"]["line_spacing"]),
            "space_before_pt": compact_counts(selected, lambda item: item["paragraph"]["space_before_pt"]),
            "space_after_pt": compact_counts(selected, lambda item: item["paragraph"]["space_after_pt"]),
            "examples": selected[:10],
        }
    return {
        "file": str(path.resolve()),
        "paragraphs": len(paragraphs),
        "tables": len(table_records),
        "inline_shapes": len(document.inline_shapes),
        "sections": sections,
        "roles": summaries,
        "table_profiles": table_records,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("files", nargs="+", type=Path)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()
    result = [extract(path) for path in args.files]
    payload: Any = result[0] if len(result) == 1 else result
    output = json.dumps(payload, ensure_ascii=False, indent=2) + "\n"
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(output, encoding="utf-8")
        print(args.out.resolve())
    else:
        print(output, end="")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
