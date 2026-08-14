#!/usr/bin/env python3
"""Create a deterministic project manifest, folder tree, attention file, index, and workbook scaffold."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import sys
from datetime import datetime
from pathlib import Path

from apply_docx_format_contract import apply_in_place
from build_generic_docx_templates import set_repeat_header, set_table_geometry, setup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from generate_attention_items import build_document as build_attention_document
from manual_acceptance import pending_acceptance
from project_blueprint import FOLDERS, materials_for_scope
from validate_project_manifest import validate

REQUIRED_INTAKE_PATHS = (
    ("project", "title"),
    ("project", "leader"),
    ("project", "school"),
    ("project", "subject"),
    ("project", "stage"),
    ("project", "year"),
    ("governance", "current_date"),
    ("governance", "project_status"),
    ("generation_contract", "package_scope"),
    ("generation_contract", "truth_state"),
    ("submission_requirements", "authority"),
    ("timeline", "application"),
    ("timeline", "completion"),
)

DELIVERY_BY_SCOPE = {
    "application-kit": "application-scaffold",
    "implementation-kit": "implementation-scaffold",
    "full-lifecycle-kit": "full-lifecycle-scaffold",
    "closing-kit": "closing-scaffold",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def nested_value(data: dict, path: tuple[str, str]) -> object:
    parent = data.get(path[0])
    return parent.get(path[1]) if isinstance(parent, dict) else None


def build_manifest(intake: dict) -> dict:
    missing = [".".join(path) for path in REQUIRED_INTAKE_PATHS if nested_value(intake, path) in (None, "", [])]
    if missing:
        raise ValueError(f"初始化输入缺少关键字段：{', '.join(missing)}")

    scope = str(intake["generation_contract"]["package_scope"])
    if scope not in DELIVERY_BY_SCOPE:
        raise ValueError(f"初始化器不支持package_scope={scope!r}；请使用四种标准成套范围")
    current_date = str(intake["governance"]["current_date"])
    generation = dict(intake["generation_contract"])
    generation.setdefault("delivery_state", DELIVERY_BY_SCOPE[scope])
    generation.setdefault("batch_mode", "single-snapshot")
    generation.setdefault("unknown_handling", "structured-pending")
    generation.setdefault("target_versions", ["working", "submission"])
    generation.setdefault("coverage_exemptions", [])
    generation.setdefault("snapshot_id", f"snapshot-{current_date}-01")
    generation.setdefault("parent_snapshot_id", None)
    generation.setdefault("generated_at", f"{current_date}T00:00:00+08:00")

    requirements = dict(intake["submission_requirements"])
    requirements.setdefault("status", "pending")
    requirements.setdefault("year", intake["project"]["year"])
    requirements.setdefault("verified_at", None)
    requirements.setdefault("search_run_id", None)
    requirements.setdefault("searched_at", None)
    requirements.setdefault("official_portals_checked", [])
    requirements.setdefault("search_queries", [])
    requirements.setdefault("policy_snapshot_file", None)
    requirements.setdefault("policy_snapshot_sha256", None)
    requirements.setdefault("deadline", None)
    requirements.setdefault("notice_source_ids", [])
    requirements.setdefault("template_source_ids", [])
    requirements.setdefault("required_material_ids", [])
    requirements.setdefault("anonymous_required", None)
    requirements.setdefault("submission_mode", "pending")
    requirements.setdefault("file_rules", {"max_size_mb": None, "naming_rule": None, "copies": None})

    manifest = {
        "schema_version": "1.6",
        "project": intake["project"],
        "project_context": intake.get("project_context", {}),
        "problem_context": intake.get("problem_context", {}),
        "governance": intake["governance"],
        "generation_contract": generation,
        "manual_acceptance": pending_acceptance(),
        "submission_requirements": requirements,
        "contributors": intake.get("contributors", []),
        "subject_coverage": intake.get("subject_coverage", []),
        "timeline": intake["timeline"],
        "research_questions": intake.get("research_questions", []),
        "logic_mappings": intake.get("logic_mappings", []),
        "samples": intake.get("samples", []),
        "instruments": intake.get("instruments", []),
        "evidence": intake.get("evidence", []),
        "interventions": intake.get("interventions", []),
        "cases": intake.get("cases", []),
        "sources": intake.get("sources", []),
        "commitments": intake.get("commitments", []),
        "claims": intake.get("claims", []),
        "evaluation_weights": intake.get("evaluation_weights", {}),
        "legacy_forbidden_terms": intake.get("legacy_forbidden_terms", []),
        "materials": materials_for_scope(scope),
    }
    errors, _ = validate(manifest)
    if errors:
        raise ValueError("初始化输入未通过主清单校验：\n- " + "\n- ".join(errors))
    return manifest


def build_index(manifest: dict, target: Path) -> None:
    document = Document()
    setup(document, "课题材料目录与交付索引")
    section = document.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("课题材料目录与交付索引").bold = True
    document.add_paragraph(f"课题：{manifest['project']['title']}")
    document.add_paragraph(f"主清单快照：{manifest['generation_contract']['snapshot_id']}")
    table = document.add_table(rows=1, cols=8)
    set_table_geometry(table, [700, 2200, 1500, 1100, 1100, 900, 3900, 2300])
    set_repeat_header(table.rows[0])
    headers = ("ID", "材料名称", "角色", "阶段", "状态", "格式", "计划文件", "依赖")
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for item in manifest["materials"]:
        dependencies = item.get("depends_on", [])
        dependency_text = "、".join(dependencies)
        if item.get("material_role") == "index" and len(dependencies) > 5:
            dependency_text = f"本批全部材料（{len(dependencies)}项）"
        row = table.add_row().cells
        values = (
            item["id"], item["name"], item["material_role"], item["stage"], item["status"],
            item["output_format"], item.get("planned_file_path") or "", dependency_text,
        )
        for cell, value in zip(row, values):
            cell.text = str(value)
    document.add_paragraph("说明：本索引由项目主清单自动生成。正式交付前须刷新状态、文件路径、哈希和QA记录。")
    document.save(target)


def sync_index_row(manifest: dict, target: Path) -> None:
    """Keep the generated index's own row consistent without recursive self-hashing."""
    document = Document(target)
    if not document.tables:
        raise ValueError("交付索引缺少材料表")
    table = document.tables[0]
    item = next(value for value in manifest["materials"] if value["id"] == "M25")
    for row in table.rows[1:]:
        if row.cells[0].text.strip() != "M25":
            continue
        values = (
            item["id"], item["name"], item["material_role"], item["stage"], item["status"],
            item["output_format"], item.get("file_path") or item.get("planned_file_path") or "",
            f"本批全部材料（{len(item.get('depends_on', []))}项）",
        )
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
        break
    document.save(target)


def initialize(intake_path: Path, root: Path, skill_root: Path) -> Path:
    intake = json.loads(intake_path.read_text(encoding="utf-8"))
    manifest = build_manifest(intake)
    root = root.resolve()
    manifest_path = root / "project-manifest.json"
    controlled = (manifest_path, root / "00_课题材料包注意事项_真实性与待办清单_v0.1.docx", root / "M25_整套材料目录与交付索引_v0.1.docx")
    conflicts = [str(path) for path in controlled if path.exists()]
    if conflicts:
        raise FileExistsError("拒绝覆盖已有项目控制文件：" + "，".join(conflicts))

    root.mkdir(parents=True, exist_ok=True)
    for folder in FOLDERS:
        (root / folder).mkdir(parents=True, exist_ok=True)

    material_by_id = {item["id"]: item for item in manifest["materials"]}
    attention_target = root / "00_课题材料包注意事项_真实性与待办清单_v0.1.docx"
    index_target = root / "M25_整套材料目录与交付索引_v0.1.docx"
    material_by_id["M00"].update(status="draft", file_path=attention_target.name)
    material_by_id["M25"].update(status="draft", file_path=index_target.name)
    workbook_item = material_by_id.get("M09")
    if workbook_item and workbook_item.get("included_in_batch") is True:
        workbook_target = root / str(workbook_item["planned_file_path"])
        workbook_target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(skill_root / "assets" / "templates" / "project-data-workbook.xlsx", workbook_target)
        workbook_item.update(status="draft", file_path=str(workbook_target.relative_to(root)), sha256=sha256(workbook_target))

    build_index(manifest, index_target)
    sync_index_row(manifest, index_target)
    apply_in_place(index_target, "M25")
    material_by_id["M25"]["sha256"] = sha256(index_target)

    from datetime import date

    attention_day = date.fromisoformat(str(manifest["governance"]["current_date"]))
    build_attention_document(manifest, attention_target, attention_day)
    apply_in_place(attention_target, "M00")
    material_by_id["M00"]["sha256"] = sha256(attention_target)

    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--intake", type=Path, required=True, help="一次采集信息JSON")
    parser.add_argument("--root", type=Path, required=True, help="新材料包根目录")
    parser.add_argument("--skill-root", type=Path, default=Path(__file__).resolve().parents[1])
    args = parser.parse_args()
    try:
        manifest_path = initialize(args.intake, args.root, args.skill_root.resolve())
    except (OSError, json.JSONDecodeError, ValueError) as exc:
        print(f"初始化失败：{exc}", file=sys.stderr)
        return 1
    print(f"初始化完成：{manifest_path}")
    print(f"生成时间：{datetime.now().astimezone().isoformat(timespec='seconds')}")
    print("注意：生成的是可继续制作的脚手架；完成逐页/逐表QA前，自动生成文件保持draft。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
