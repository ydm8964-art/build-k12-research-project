#!/usr/bin/env python3
"""Structural DOCX format audit for K-12 research-project materials."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from audit_common import PLACEHOLDER_RE, cli_failed
from docx import Document
from docx.oxml.ns import qn

A4_WIDTH_MM = 210.0
A4_HEIGHT_MM = 297.0
TOLERANCE_MM = 1.0
TABLE_PROFILES = {"lesson-table", "casebook", "attention-items"}
REPEAT_HEADER_PROFILES = {"research-form", "analysis-report", "casebook", "evidence-sheet", "attention-items"}
VALID_PROFILES = {
    "official-exact",
    "research-form",
    "analysis-report",
    "lesson-table",
    "lesson-long",
    "casebook",
    "evidence-sheet",
    "attention-items",
}


def close(left: float, right: float, tolerance: float = TOLERANCE_MM) -> bool:
    return abs(left - right) <= tolerance


def section_signature(section) -> tuple[float, ...]:
    return tuple(
        round(value.mm, 1)
        for value in (
            section.page_width,
            section.page_height,
            section.top_margin,
            section.bottom_margin,
            section.left_margin,
            section.right_margin,
            section.header_distance,
            section.footer_distance,
        )
    ) + (int(section.orientation),)


def table_width_dxa(table) -> int | None:
    node = table._tbl.tblPr.find(qn("w:tblW"))
    if node is None or node.get(qn("w:type")) != "dxa":
        return None
    try:
        return int(node.get(qn("w:w")))
    except (TypeError, ValueError):
        return None


def integer_attr(node, name: str) -> int | None:
    if node is None:
        return None
    try:
        return int(node.get(qn(name)))
    except (TypeError, ValueError):
        return None


def body_block_signature(document: Document) -> tuple[str, ...]:
    return tuple(child.tag.rsplit("}", 1)[-1] for child in document.element.body.iterchildren())


def document_feature_signature(document: Document) -> tuple[int, int, int, int]:
    root = document.element
    return (
        len(root.xpath(".//w:sdt")),
        len(root.xpath(".//w:fldSimple")),
        len(root.xpath(".//w:instrText")),
        len(root.xpath(".//w:drawing|.//w:pict")),
    )


def table_structure_signature(table) -> tuple:
    grid = tuple(integer_attr(col, "w:w") for col in table._tbl.tblGrid.findall(qn("w:gridCol")))
    cells: list[tuple[int, str, int | None]] = []
    for row_index, row in enumerate(table._tbl.findall(qn("w:tr")), 1):
        for cell in row.findall(qn("w:tc")):
            props = cell.find(qn("w:tcPr"))
            span_node = props.find(qn("w:gridSpan")) if props is not None else None
            merge_node = props.find(qn("w:vMerge")) if props is not None else None
            span = integer_attr(span_node, "w:val") or 1
            merge = merge_node.get(qn("w:val"), "continue") if merge_node is not None else ""
            cells.append((row_index, merge, span))
    return grid, tuple(cells)


def has_repeat_header(table) -> bool:
    if not table.rows:
        return False
    props = table.rows[0]._tr.trPr
    return props is not None and props.find(qn("w:tblHeader")) is not None


def has_cell_margins(table) -> bool:
    return table._tbl.tblPr.find(qn("w:tblCellMar")) is not None


def exact_height_rows(table) -> list[int]:
    found: list[int] = []
    for index, row in enumerate(table.rows, 1):
        props = row._tr.trPr
        if props is None:
            continue
        for height in props.findall(qn("w:trHeight")):
            if height.get(qn("w:hRule")) == "exact":
                found.append(index)
                break
    return found


def all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def audit(
    path: Path,
    profile: str,
    reference: Path | None,
    final: bool,
    allow_added_drawings: bool = False,
) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    if profile not in VALID_PROFILES:
        return [f"未知profile：{profile}"], warnings
    if path.suffix.lower() != ".docx":
        return ["结构审计仅支持DOCX；DOC请先保留原件并制作工作副本"], warnings

    document = Document(path)
    if not document.sections:
        errors.append("文档没有分节信息")

    if profile == "official-exact":
        if reference is None:
            errors.append("official-exact必须提供--reference")
        else:
            source = Document(reference)
            if len(document.sections) != len(source.sections):
                errors.append(f"分节数量与官方模板不一致：{len(document.sections)} != {len(source.sections)}")
            for index, (actual, expected) in enumerate(zip(document.sections, source.sections), 1):
                if section_signature(actual) != section_signature(expected):
                    errors.append(f"第{index}节纸张/页边距/页眉页脚与官方模板不一致")
            if len(document.tables) != len(source.tables):
                errors.append(f"表格数量与官方模板不一致：{len(document.tables)} != {len(source.tables)}")
            for index, (actual, expected) in enumerate(zip(document.tables, source.tables), 1):
                actual_shape = (len(actual.rows), len(actual.columns))
                expected_shape = (len(expected.rows), len(expected.columns))
                if actual_shape != expected_shape:
                    errors.append(f"第{index}个表格形状与模板不一致：{actual_shape} != {expected_shape}")
                if table_structure_signature(actual) != table_structure_signature(expected):
                    errors.append(f"第{index}个表格的列宽网格或合并单元格结构与官方模板不一致")
            if body_block_signature(document) != body_block_signature(source):
                errors.append("正文段落/表格/分节的块级顺序与官方模板不一致")
            actual_features = document_feature_signature(document)
            source_features = document_feature_signature(source)
            if actual_features[:3] != source_features[:3]:
                errors.append("内容控件或域数量与官方模板不一致")
            if allow_added_drawings:
                if actual_features[3] < source_features[3]:
                    errors.append("官方模板原有绘图或图片被删除")
            elif actual_features[3] != source_features[3]:
                errors.append("绘图或图片数量与官方模板不一致；真实签章/照片确需新增时登记allow_added_drawings")
    else:
        for index, section in enumerate(document.sections, 1):
            width, height = section.page_width.mm, section.page_height.mm
            portrait_a4 = close(width, A4_WIDTH_MM) and close(height, A4_HEIGHT_MM)
            landscape_a4 = close(width, A4_HEIGHT_MM) and close(height, A4_WIDTH_MM)
            if not (portrait_a4 or landscape_a4):
                errors.append(f"第{index}节不是A4：{width:.1f}×{height:.1f}mm")
            if landscape_a4 and profile not in {"analysis-report", "evidence-sheet", "attention-items"}:
                warnings.append(f"第{index}节为横向A4，请确认该材料确需横向")
            if min(section.top_margin.mm, section.bottom_margin.mm, section.left_margin.mm, section.right_margin.mm) < 15:
                warnings.append(f"第{index}节页边距小于15mm，可能不适合装订或打印")

    if profile in TABLE_PROFILES and not document.tables:
        errors.append(f"{profile}应至少包含一个表格")

    for index, table in enumerate(document.tables, 1):
        section = document.sections[0]
        usable_mm = section.page_width.mm - section.left_margin.mm - section.right_margin.mm
        width = table_width_dxa(table)
        if width is None or width == 0:
            warnings.append(f"第{index}个表格未设置明确DXA总宽")
        else:
            width_mm = width / 1440 * 25.4
            if width_mm > usable_mm + 1.0:
                errors.append(f"第{index}个表格宽{width_mm:.1f}mm，超过版心{usable_mm:.1f}mm")
        grid = table._tbl.tblGrid.findall(qn("w:gridCol"))
        if not grid or any(not col.get(qn("w:w")) for col in grid):
            warnings.append(f"第{index}个表格缺少完整列宽网格")
            grid_widths: list[int] = []
        else:
            grid_widths = [int(col.get(qn("w:w"))) for col in grid]
            if width and abs(sum(grid_widths) - width) > 10:
                warnings.append(f"第{index}个表格tblGrid合计与tblW不一致")
        indent = table._tbl.tblPr.find(qn("w:tblInd"))
        if integer_attr(indent, "w:w") is None or indent.get(qn("w:type")) != "dxa":
            warnings.append(f"第{index}个表格未设置明确DXA缩进tblInd")
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        if layout is None or layout.get(qn("w:type")) != "fixed":
            warnings.append(f"第{index}个表格未使用fixed布局，跨软件列宽可能漂移")
        if grid_widths:
            for row_number, row in enumerate(table._tbl.findall(qn("w:tr")), 1):
                grid_index = 0
                for cell_number, cell in enumerate(row.findall(qn("w:tc")), 1):
                    props = cell.find(qn("w:tcPr"))
                    span_node = props.find(qn("w:gridSpan")) if props is not None else None
                    span = integer_attr(span_node, "w:val") or 1
                    cell_width_node = props.find(qn("w:tcW")) if props is not None else None
                    cell_width = integer_attr(cell_width_node, "w:w")
                    expected_width = sum(grid_widths[grid_index : grid_index + span])
                    if cell_width is None or cell_width_node.get(qn("w:type")) != "dxa":
                        warnings.append(f"第{index}个表格第{row_number}行第{cell_number}格缺少明确DXA tcW")
                    elif expected_width and abs(cell_width - expected_width) > 10:
                        warnings.append(f"第{index}个表格第{row_number}行第{cell_number}格tcW与跨列网格不一致")
                    grid_index += span
        if len(table.rows) > 2 and profile in REPEAT_HEADER_PROFILES and not has_repeat_header(table):
            warnings.append(f"第{index}个多行表格未设置重复表头")
        if not has_cell_margins(table):
            warnings.append(f"第{index}个表格未设置表级单元格内边距")
        exact_rows = exact_height_rows(table)
        if exact_rows:
            warnings.append(f"第{index}个表格存在固定行高，可能截字：{exact_rows[:8]}")

    text = all_text(document)
    if profile == "casebook" and "目录" not in text:
        errors.append("案例集缺少目录")
    if final:
        count = len(PLACEHOLDER_RE.findall(text))
        if count:
            warnings.append(f"发现结构化占位内容{count}处；仅官方保留的签章/线下填写区可以存在")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.text.strip() and run.font.size and run.font.size.pt < 9:
                warnings.append("正文存在小于9磅文字，请检查可读性")
                return errors, warnings
    return errors, warnings


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--profile", required=True, choices=sorted(VALID_PROFILES))
    parser.add_argument("--reference", type=Path)
    parser.add_argument("--allow-added-drawings", action="store_true", help="允许在保留官方原图的前提下新增真实签章/照片")
    parser.add_argument("--final", action="store_true")
    args = parser.parse_args()
    try:
        errors, warnings = audit(args.docx, args.profile, args.reference, args.final, args.allow_added_drawings)
    except Exception as exc:  # keep CLI failure legible for production use
        print(f"读取或审计失败：{exc}", file=sys.stderr)
        return 2
    for item in warnings:
        print(f"警告：{item}")
    for item in errors:
        print(f"错误：{item}")
    if cli_failed(errors, warnings, args.final):
        print(f"格式结构审计未通过：{len(errors)}个错误，{len(warnings)}个警告")
        return 1
    print(f"格式结构审计通过：0个错误，{len(warnings)}个警告")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
