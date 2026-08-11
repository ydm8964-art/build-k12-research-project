#!/usr/bin/env python3
"""Build sanitized generic DOCX assets for non-official K-12 research materials."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

# Keep the structural assets portable for Word/WPS on common school computers.
# Final documents replace these with the fonts required by the annual template.
BODY_FONT = "宋体"
HEADING_FONT = "黑体"
CONTENT_WIDTH_DXA = 8736  # A4, left/right 28 mm


def set_run(run, font: str = BODY_FONT, size: float = 12, bold: bool = False) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    props = run._element.get_or_add_rPr()
    fonts = props.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), font)


def setup(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(28)
    section.right_margin = Mm(28)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    for style_name, size in (("Title", 22), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = HEADING_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
    doc.core_properties.title = title
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    add_page_field(section.footer.paragraphs[0])


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run(run, size=10.5)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    run2 = paragraph.add_run(" 页")
    set_run(run2, size=10.5)


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    # Do not inherit the built-in Title style: some Word/LibreOffice defaults add
    # a themed rule below it, which is not part of the source material format.
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    set_run(paragraph.add_run(text), HEADING_FONT, 22, True)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        set_run(p.add_run(subtitle), BODY_FONT, 12)


def set_repeat_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    props.append(element)


def set_row_min_height(row, height_mm: float) -> None:
    row.height = Mm(height_mm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_cell_margins(table, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    props = table._tbl.tblPr
    old = props.find(qn("w:tblCellMar"))
    if old is not None:
        props.remove(old)
    margins = OxmlElement("w:tblCellMar")
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    props.append(margins)


def set_table_geometry(table, widths: list[int]) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    props = table._tbl.tblPr
    old_borders = props.find(qn("w:tblBorders"))
    if old_borders is not None:
        props.remove(old_borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), "808080")
        borders.append(node)
    props.append(borders)
    width = props.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        props.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = props.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell.width = Mm(value / 1440 * 25.4)
            props_cell = cell._tc.get_or_add_tcPr()
            tcw = props_cell.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                props_cell.append(tcw)
            tcw.set(qn("w:w"), str(value))
            tcw.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(table)


def set_three_line_table(table) -> None:
    """Replace grid borders with a top line, header rule, and bottom line."""
    props = table._tbl.tblPr
    old = props.find(qn("w:tblBorders"))
    if old is not None:
        props.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge, value, size in (
        ("top", "single", "8"), ("left", "nil", "0"),
        ("bottom", "single", "8"), ("right", "nil", "0"),
        ("insideH", "nil", "0"), ("insideV", "nil", "0"),
    ):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), "000000")
        borders.append(node)
    props.append(borders)
    for cell in table.rows[0].cells:
        cell_props = cell._tc.get_or_add_tcPr()
        cell_borders = cell_props.find(qn("w:tcBorders"))
        if cell_borders is None:
            cell_borders = OxmlElement("w:tcBorders")
            cell_props.append(cell_borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "000000")
        cell_borders.append(bottom)


def fill_cell(cell, text: str, bold: bool = False, center: bool = False, size: float = 10.5) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.clear()
    set_run(paragraph.add_run(text), BODY_FONT, size, bold)


def shade(cell, fill: str = "E7E6E6") -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), HEADING_FONT, 16 if level == 1 else 14, True)


def build_research_form(path: Path) -> None:
    doc = Document(); setup(doc, "研究工具通用母版")
    add_title(doc, "【工具名称】", "【课题规范题目】")
    p = doc.add_paragraph("填写说明：【用途、对象、匿名/知情说明、预计用时和填写方法】")
    p.paragraph_format.space_after = Pt(8)
    heading(doc, "一、基本信息")
    table = doc.add_table(rows=3, cols=4); set_table_geometry(table, [1500, 2868, 1500, 2868])
    set_repeat_header(table.rows[0])
    for row, values in zip(table.rows, (("记录ID", "【编码】", "日期", "【yyyy-mm-dd】"), ("对象/班级", "【填写】", "地点", "【填写】"), ("实施/记录人", "【填写】", "版本", "【填写】"))):
        for i, value in enumerate(values): fill_cell(row.cells[i], value, bold=i % 2 == 0, center=i % 2 == 0)
    heading(doc, "二、调查/观察项目")
    scale = doc.add_table(rows=6, cols=6); set_table_geometry(scale, [4236, 900, 900, 900, 900, 900]); set_repeat_header(scale.rows[0])
    headers = ("项目/可观察指标", "非常符合", "比较符合", "一般", "不太符合", "完全不符合")
    for i, value in enumerate(headers): fill_cell(scale.rows[0].cells[i], value, True, True, 10); shade(scale.rows[0].cells[i])
    for r in range(1, 6):
        fill_cell(scale.rows[r].cells[0], f"{r}.【根据指标矩阵生成的题项】")
        for c in range(1, 6): fill_cell(scale.rows[r].cells[c], "□", center=True)
    heading(doc, "三、开放记录")
    for label in ("1.【开放问题/关键事件】", "2.【补充说明/改进建议】"):
        doc.add_paragraph(label)
        doc.add_paragraph("\n\n")
    doc.save(path)


def build_analysis_report(path: Path) -> None:
    doc = Document(); setup(doc, "分析报告通用母版")
    add_title(doc, "【调查/访谈/观察/测评】分析报告", "【课题规范题目】")
    for title in ("一、分析目的", "二、对象、样本与工具", "三、实施过程与数据处理", "四、分维度结果"):
        heading(doc, title); doc.add_paragraph("【从项目主清单和真实数据生成；不得填写计划样本作为实际样本】")
    p = doc.add_paragraph("表1 【结果表名称】（N=【实际有效样本】）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=5, cols=4); set_table_geometry(table, [3300, 1800, 1800, 1836]); set_repeat_header(table.rows[0]); set_three_line_table(table)
    for i, value in enumerate(("项目/选项", "频数", "比例", "统计口径")): fill_cell(table.rows[0].cells[i], value, True, True)
    for r in range(1, 5):
        for c, value in enumerate((f"【项目{r}】", "【公式结果】", "【0.0%】", "【分母/多选说明】")): fill_cell(table.rows[r].cells[c], value, center=c in {1, 2})
    doc.add_paragraph("注：数据来源为《【工作簿名】》“【工作表】”【区域】；数据截止日期【yyyy-mm-dd】。")
    for title in ("五、交叉/对比与质性主题", "六、多源证据互证", "七、主要结论与干预启示", "八、局限"):
        heading(doc, title); doc.add_paragraph("【基于真实数据、记录ID和证据ID生成】")
    doc.save(path)


def build_lesson_table(path: Path) -> None:
    doc = Document(); setup(doc, "常态课堂教学设计母版")
    add_title(doc, "【教材版本·年级·册次】《【课题】》教学设计")
    table = doc.add_table(rows=15, cols=4); set_table_geometry(table, [1400, 2968, 2968, 1400])
    rows = [
        ("授课教师", "【姓名】", "学校", "【学校全称】"), ("教学课题", "【课题】", "教材版本", "【版本】"),
        ("教学对象", "【年级班级】", "课时", "【课时】"), ("课标依据", "【课标内容与行为动词】", "", ""),
        ("教材分析", "【内容结构和地位】", "", ""), ("学情分析", "【已有基础、困难和差异】", "", ""),
        ("素养目标", "【可观察目标】", "", ""), ("重点难点", "【重点/难点及依据】", "", ""),
        ("方法资源", "【教法、学法、资源与来源】", "", ""),
        ("教学环节", "教师活动", "学生活动", "设计意图/评价"),
        ("导入", "【情境与问题】", "【观察与回应】", "【诊断】"),
        ("任务一", "【支架与组织】", "【具体活动与产出】", "【标准与证据】"),
        ("跨学科微环节", "【真实问题；8—10分钟；关联学科方法】", "【任务、分工和成果】", "【量规/反馈】"),
        ("总结作业", "【小结与分层作业】", "【成果提交】", "【评价】"),
        ("板书与反思", "【板书结构】", "【实施后填写反思】", "【证据ID/改进】"),
    ]
    for r, values in enumerate(rows):
        for c, value in enumerate(values): fill_cell(table.rows[r].cells[c], value, bold=(r == 9 or c == 0), center=(r == 9 or c == 0), size=10)
        if r == 9:
            set_repeat_header(table.rows[r])
            for cell in table.rows[r].cells: shade(cell)
    doc.save(path)


def build_lesson_long(path: Path) -> None:
    doc = Document(); setup(doc, "示范课教学设计母版")
    add_title(doc, "《【课题】》教学设计", "【教材版本·年级·课时｜教师｜学校】")
    doc.add_paragraph("目录（完成正文后在Word中更新自动目录）")
    sections = (
        "一、教学理论依据与设计理念", "二、课程标准分析", "三、教材分析", "四、学情分析",
        "五、教学思路与问题链", "六、核心素养目标", "七、教学重难点", "八、方法、资源与准备",
    )
    for title in sections:
        heading(doc, title); doc.add_paragraph("【填写经核验的内容；理论、课标、数据和资源登记来源ID】")
    heading(doc, "九、教学过程")
    table = doc.add_table(rows=7, cols=4); set_table_geometry(table, [1500, 3000, 2500, 1736]); set_repeat_header(table.rows[0])
    for i, value in enumerate(("教学环节", "教师活动", "学生活动", "设计意图/评价")): fill_cell(table.rows[0].cells[i], value, True, True); shade(table.rows[0].cells[i])
    for r, label in enumerate(("情境导入", "任务一", "任务二", "跨学科任务", "成果展示评价", "总结迁移"), 1):
        for c, value in enumerate((label, "【教师活动、资源来源】", "【学生行动和成果】", "【评价标准和证据】")): fill_cell(table.rows[r].cells[c], value, center=c == 0)
    for title in ("十、板书设计", "十一、作业设计", "十二、教学反思", "参考文献"):
        heading(doc, title); doc.add_paragraph("【实施后据实填写；引用必须可核验】")
    doc.save(path)


def build_casebook(path: Path) -> None:
    doc = Document(); setup(doc, "教学案例集母版")
    add_title(doc, "【课题简称】教学案例集", "【负责人｜责任单位｜版本和日期】")
    heading(doc, "目录"); doc.add_paragraph("【使用Word自动目录；终稿更新案例标题和页码】")
    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "案例一：【案例标题】")
    table = doc.add_table(rows=6, cols=2); set_table_geometry(table, [1800, 6936])
    set_repeat_header(table.rows[0])
    for row, values in zip(table.rows, (("案例ID", "【CASE-…】"), ("状态", "【designed/piloted/implemented/validated】"), ("作者/贡献", "【作者ID、贡献和授权】"), ("实施信息", "【日期、学校、班级、教师、人数、课时、版本】"), ("证据", "【照片/作品/观察/评价ID】"), ("安全与审批", "【适用时填写】"))):
        fill_cell(row.cells[0], values[0], True, True); fill_cell(row.cells[1], values[1])
    for title in ("一、背景与问题", "二、学情与目标", "三、跨学科内容与核心任务", "四、准备、材料与安全"):
        heading(doc, title); doc.add_paragraph("【按案例状态填写；开发稿只能写预期，实施稿必须引用真实证据ID】")
    doc.add_page_break()
    heading(doc, "五、实施过程")
    table2 = doc.add_table(rows=5, cols=4); set_table_geometry(table2, [1500, 2800, 2500, 1936]); set_repeat_header(table2.rows[0])
    for i, value in enumerate(("环节", "教师活动", "学生活动", "成果/证据")): fill_cell(table2.rows[0].cells[i], value, True, True); shade(table2.rows[0].cells[i])
    for r in range(1, 5):
        for c, value in enumerate((f"【环节{r}】", "【填写】", "【填写】", "【证据ID】")): fill_cell(table2.rows[r].cells[c], value, center=c == 0)
    for title in ("六、任务单、作品与评价", "七、结果、反思与改进"):
        heading(doc, title); doc.add_paragraph("【按案例状态填写；开发稿只能写预期，实施稿必须引用真实证据ID】")
    doc.save(path)


def build_evidence_sheet(path: Path) -> None:
    doc = Document(); setup(doc, "照片证据册母版")
    add_title(doc, "【课题简称】研究过程照片证据册", "【隐私级别｜版本｜数据截止日期】")
    heading(doc, "一、照片索引")
    index = doc.add_table(rows=5, cols=6); set_table_geometry(index, [700, 1250, 1200, 2100, 1600, 1886]); set_repeat_header(index.rows[0])
    for i, value in enumerate(("图号", "照片ID", "日期", "活动", "对应材料/案例", "授权/人物处理")): fill_cell(index.rows[0].cells[i], value, True, True, 9.5); shade(index.rows[0].cells[i])
    for r in range(1, 5):
        for c, value in enumerate((f"图{r}", "【PHO-…】", "【日期】", "【活动】", "【材料/案例ID】", "【状态】")): fill_cell(index.rows[r].cells[c], value, center=c in {0, 1, 2}, size=9.5)
    doc.add_page_break()
    heading(doc, "二、分阶段照片记录")
    for r in range(1, 3):
        frame = doc.add_table(rows=2, cols=1); set_table_geometry(frame, [CONTENT_WIDTH_DXA])
        set_row_min_height(frame.rows[0], 58)
        fill_cell(frame.rows[0].cells[0], f"【在此内嵌真实照片 PHO-年份-{r:03d}；保持比例；不得使用网络图或生成图】", center=True)
        fill_cell(frame.rows[1].cells[0], f"图{r} 【活动名称】（证据ID：【PHO-…】；日期：【yyyy-mm-dd】；地点：【填写】）", center=True)
        doc.add_paragraph()
    heading(doc, "三、原件与授权说明")
    doc.add_paragraph("【登记原件文件名、SHA-256、拍摄来源、变换记录、授权范围和受控原件位置】")
    doc.save(path)


def build_attention_items(path: Path) -> None:
    doc = Document(); setup(doc, "课题材料包注意事项与真实性待办清单")
    section = doc.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    add_title(doc, "【课题简称】材料包注意事项", "真实性、缺件、时间逻辑与增强建议｜【批次日期】")

    note = doc.add_table(rows=1, cols=1); set_table_geometry(note, [14520])
    fill_cell(
        note.rows[0].cells[0],
        "本文件供课题负责人内部核对，不代替官方申报/结题附件。所有照片、数据、原始记录、证明、签字和盖章必须来自真实项目；本文件随材料包每次重建。",
        size=10.5,
    )
    shade(note.rows[0].cells[0], "FFF2CC")

    heading(doc, "一、材料包状态快照")
    overview = doc.add_table(rows=3, cols=4); set_table_geometry(overview, [1900, 5360, 1900, 5360])
    set_repeat_header(overview.rows[0])
    values = (
        ("课题/负责人", "【由项目主清单生成】", "单位/阶段", "【由项目主清单生成】"),
        ("成套范围", "【package_scope/真实性阶段】", "目标状态", "【delivery_state】"),
        ("统计", "【阻断/必须/阶段/建议数量】", "数据/要求快照", "【截止日期/核验状态】"),
    )
    for r, row_values in enumerate(values):
        for c, value in enumerate(row_values):
            fill_cell(overview.rows[r].cells[c], value, bold=c % 2 == 0, center=c % 2 == 0, size=10)

    heading(doc, "二、缺失、矛盾与待人工完成事项")
    missing = doc.add_table(rows=4, cols=7); set_table_geometry(missing, [650, 900, 1300, 3000, 2500, 3600, 2570]); set_repeat_header(missing.rows[0])
    headers = ("ID", "级别", "类别", "具体事项", "为什么需要", "建议行动/责任/最迟时间", "对应材料与完成标准")
    for i, value in enumerate(headers): fill_cell(missing.rows[0].cells[i], value, True, True, 9); shade(missing.rows[0].cells[i])
    for r in range(1, 4):
        row_values = (f"ATT-{r:03d}", "【级别】", "【类别】", "【从主清单生成】", "【证据或格式原因】", "【行动、责任人、日期】", "【材料ID/验收标准】")
        for c, value in enumerate(row_values): fill_cell(missing.rows[r].cells[c], value, center=c in {0, 1}, size=9)

    heading(doc, "三、真实照片拍摄、提供与插入清单")
    photos = doc.add_table(rows=4, cols=7); set_table_geometry(photos, [850, 1200, 2600, 2800, 2200, 2700, 2170]); set_repeat_header(photos.rows[0])
    headers = ("照片ID", "阶段", "真实活动/建议画面", "必须同步取得的信息", "授权/人物处理", "目标材料与插入位置", "最迟时间/状态")
    for i, value in enumerate(headers): fill_cell(photos.rows[0].cells[i], value, True, True, 9); shade(photos.rows[0].cells[i])
    for r in range(1, 4):
        row_values = (f"PHO-年份-{r:03d}", "【阶段】", "【真实活动，不得摆拍冒充】", "【日期、地点、来源、记录ID】", "【授权和打码】", "【材料ID/章节】", "【日期/状态】")
        for c, value in enumerate(row_values): fill_cell(photos.rows[r].cells[c], value, center=c in {0, 1}, size=9)

    heading(doc, "四、其他真实原始材料清单")
    originals = doc.add_table(rows=4, cols=6); set_table_geometry(originals, [1200, 2200, 3400, 2400, 2700, 2620]); set_repeat_header(originals.rows[0])
    headers = ("类别", "真实原件", "最低标准", "命名/保存目录", "对应时间与用途", "状态/责任人")
    for i, value in enumerate(headers): fill_cell(originals.rows[0].cells[i], value, True, True, 9); shade(originals.rows[0].cells[i])
    for r in range(1, 4):
        row_values = ("【类别】", "【原件名称】", "【页序、来源、签章或数据标准】", "【文件名/受控目录】", "【节点/目标材料】", "【状态/责任人】")
        for c, value in enumerate(row_values): fill_cell(originals.rows[r].cells[c], value, center=c == 0, size=9)

    heading(doc, "五、学科专项真实性与安全核验")
    subject = doc.add_table(rows=3, cols=6); set_table_geometry(subject, [1300, 2800, 2500, 2500, 2900, 2520]); set_repeat_header(subject.rows[0])
    headers = ("学科/领域", "应取得的真实材料", "学习或效果证据", "关键学科事实", "安全/伦理/版权", "对应材料与完成标准")
    for i, value in enumerate(headers): fill_cell(subject.rows[0].cells[i], value, True, True, 9); shade(subject.rows[0].cells[i])
    rows = (
        ("【主学科】", "【由全学科专项矩阵生成】", "【不能只用照片或满意度】", "【课标/教材/概念/术语核验】", "【按学科生成安全、伦理和版权要求】", "【材料ID/原件/版本/评价口径】"),
        ("【关联学科】", "【跨学科时填写；不适用写明理由】", "【相称的过程与结果证据】", "【关联学科教师或权威来源】", "【专项风险与授权】", "【核验人/日期/完成标准】"),
    )
    for r, row_values in enumerate(rows, 1):
        for c, value in enumerate(row_values): fill_cell(subject.rows[r].cells[c], value, center=c == 0, size=9)

    heading(doc, "六、时间逻辑检查")
    timeline = doc.add_table(rows=5, cols=7); set_table_geometry(timeline, [1500, 1400, 1800, 2400, 1500, 3200, 2720]); set_repeat_header(timeline.rows[0])
    headers = ("节点", "计划日期", "实际日期/证据", "前置条件", "检查结果", "处理建议", "最迟完成时间/责任人")
    for i, value in enumerate(headers): fill_cell(timeline.rows[0].cells[i], value, True, True, 9); shade(timeline.rows[0].cells[i])
    for r in range(1, 5):
        row_values = ("【研究节点】", "【yyyy-mm-dd】", "【实际日期/证据ID】", "【必须先完成事项】", "【正常/冲突/逾期】", "【不得倒改日期的处理建议】", "【日期/责任人】")
        for c, value in enumerate(row_values): fill_cell(timeline.rows[r].cells[c], value, center=c in {0, 1, 4}, size=9)

    heading(doc, "七、建议增加的增强材料")
    suggestions = doc.add_table(rows=4, cols=5); set_table_geometry(suggestions, [900, 2200, 4400, 4000, 3020]); set_repeat_header(suggestions.rows[0])
    headers = ("优先级", "建议材料", "增加理由", "建议做法/最低标准", "适用阶段/是否阻断")
    for i, value in enumerate(headers): fill_cell(suggestions.rows[0].cells[i], value, True, True, 9); shade(suggestions.rows[0].cells[i])
    for r in range(1, 4):
        row_values = ("【A/B/C】", "【建议项】", "【对应研究风险或评审价值】", "【可执行做法】", "【阶段/非阻断】")
        for c, value in enumerate(row_values): fill_cell(suggestions.rows[r].cells[c], value, center=c == 0, size=9)

    heading(doc, "八、教师最终操作顺序")
    actions = doc.add_table(rows=5, cols=5); set_table_geometry(actions, [800, 2700, 4300, 3500, 3220]); set_repeat_header(actions.rows[0])
    headers = ("顺序", "行动", "输入", "完成标准", "责任人/日期/状态")
    for i, value in enumerate(headers): fill_cell(actions.rows[0].cells[i], value, True, True, 9); shade(actions.rows[0].cells[i])
    for r in range(1, 5):
        row_values = (str(r), "【下一步行动】", "【需要取得的真实材料】", "【可核验完成标准】", "【责任人/日期/状态】")
        for c, value in enumerate(row_values): fill_cell(actions.rows[r].cells[c], value, center=c == 0, size=9)
    doc.save(path)


BUILDERS = {
    "research-form.docx": build_research_form,
    "analysis-report.docx": build_analysis_report,
    "lesson-table.docx": build_lesson_table,
    "lesson-long.docx": build_lesson_long,
    "casebook.docx": build_casebook,
    "evidence-sheet.docx": build_evidence_sheet,
    "attention-items.docx": build_attention_items,
}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    for name, builder in BUILDERS.items():
        builder(args.output_dir / name)
        print(args.output_dir / name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
