#!/usr/bin/env python3
"""Audit DOCX typography and paragraph formatting against its material contract."""

from __future__ import annotations

import argparse
import sys
import zipfile
from pathlib import Path

from apply_docx_format_contract import ALIGNMENTS, STYLE_NAMES, classify
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from format_contracts import DEFAULT_CONTRACT_PATH, material_contract


def close(actual: float | None, expected: float, tolerance: float = 0.35) -> bool:
    return actual is not None and abs(actual - expected) <= tolerance


def east_asia_font(run) -> str | None:
    rpr = run._element.rPr
    if rpr is None:
        return None
    fonts = rpr.find(qn("w:rFonts"))
    return fonts.get(qn("w:eastAsia")) if fonts is not None else None


def paragraph_xml_signature(paragraph) -> tuple[str, str | None]:
    ppr = paragraph._p.pPr
    pstyle = ppr.find(qn("w:pStyle")) if ppr is not None else None
    return (
        pstyle.get(qn("w:val")) if pstyle is not None else "",
        ppr.xml if ppr is not None else None,
    )


def first_run_xml_signature(paragraph) -> str | None:
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        return run._r.rPr.xml if run._r.rPr is not None else None
    return None


def iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def compare_official_style(reference: Path, path: Path, document: Document) -> list[str]:
    source = Document(reference)
    actual_paragraphs = list(iter_all_paragraphs(document))
    source_paragraphs = list(iter_all_paragraphs(source))
    errors: list[str] = []
    with zipfile.ZipFile(path) as actual_zip, zipfile.ZipFile(reference) as reference_zip:
        if actual_zip.read("word/styles.xml") != reference_zip.read("word/styles.xml"):
            errors.append("样式定义表与官方模板不一致；不得修改模板样式或主题字体")
    if len(actual_paragraphs) != len(source_paragraphs):
        return [f"段落及表格单元格段落数量与官方模板不一致：{len(actual_paragraphs)} != {len(source_paragraphs)}"]
    for index, (actual, expected) in enumerate(zip(actual_paragraphs, source_paragraphs), 1):
        if paragraph_xml_signature(actual) != paragraph_xml_signature(expected):
            errors.append(f"第{index}个段落的样式/缩进/行距/对齐与官方模板不一致")
        expected_run = first_run_xml_signature(expected)
        if expected_run is not None and first_run_xml_signature(actual) != expected_run:
            errors.append(f"第{index}个段落的字体/字号/加粗等字符格式与官方模板不一致")
        if len(errors) >= 30:
            errors.append("版式差异超过30处，停止继续列出")
            break
    return errors


def line_spacing_value(paragraph) -> tuple[str, float | None]:
    value = paragraph.paragraph_format.line_spacing
    if value is None:
        return "missing", None
    if hasattr(value, "pt"):
        return "fixed", round(value.pt, 2)
    multiple = round(float(value), 3)
    return ("single", multiple) if abs(multiple - 1.0) <= 0.01 else ("multiple", multiple)


def integer_attr(node, name: str) -> int | None:
    if node is None:
        return None
    try:
        return int(node.get(qn(name)))
    except (TypeError, ValueError):
        return None


def check_table_contract(table, contract: dict, label: str) -> list[str]:
    expected = contract["table"]
    errors: list[str] = []
    props = table._tbl.tblPr
    layout = props.find(qn("w:tblLayout"))
    if layout is None or layout.get(qn("w:type")) != expected["layout"]:
        errors.append(f"{label}必须使用{expected['layout']}固定布局")
    indent = props.find(qn("w:tblInd"))
    if integer_attr(indent, "w:w") != 0 or (indent is not None and indent.get(qn("w:type")) != "dxa"):
        errors.append(f"{label}表格缩进必须为0 DXA")
    margins = props.find(qn("w:tblCellMar"))
    for name, key in (
        ("top", "cell_margin_top_dxa"),
        ("bottom", "cell_margin_bottom_dxa"),
        ("start", "cell_margin_left_dxa"),
        ("end", "cell_margin_right_dxa"),
    ):
        node = margins.find(qn(f"w:{name}")) if margins is not None else None
        if integer_attr(node, "w:w") != int(expected[key]) or (node is not None and node.get(qn("w:type")) != "dxa"):
            errors.append(f"{label}{name}单元格内边距应为{expected[key]} DXA")
    if table.rows and expected.get("repeat_header"):
        row_props = table.rows[0]._tr.trPr
        if row_props is None or row_props.find(qn("w:tblHeader")) is None:
            errors.append(f"{label}首行必须设置跨页重复表头")
    if not expected.get("allow_exact_row_height"):
        for row_index, row in enumerate(table.rows, 1):
            row_props = row._tr.trPr
            if row_props is not None and any(
                node.get(qn("w:hRule")) == "exact" for node in row_props.findall(qn("w:trHeight"))
            ):
                errors.append(f"{label}第{row_index}行不得使用固定行高")
    return errors


def check_paragraph(paragraph, role_name: str, contract: dict, label: str) -> list[str]:
    expected = contract["roles"][role_name]
    errors: list[str] = []
    expected_style = STYLE_NAMES[role_name]
    if paragraph.style.name != expected_style:
        errors.append(f"{label}应使用样式{expected_style}，实际为{paragraph.style.name}")
    if paragraph.alignment != ALIGNMENTS[expected["alignment"]]:
        errors.append(f"{label}对齐应为{expected['alignment']}")
    fmt = paragraph.paragraph_format
    actual_indent = fmt.first_line_indent.pt if fmt.first_line_indent is not None else 0.0
    if not close(actual_indent, float(expected.get("first_line_indent_pt", 0.0))):
        errors.append(f"{label}首行缩进应为{expected.get('first_line_indent_pt', 0.0)}磅，实际{actual_indent:.2f}磅")
    actual_left = fmt.left_indent.pt if fmt.left_indent is not None else 0.0
    if not close(actual_left, float(expected.get("left_indent_pt", 0.0))):
        errors.append(f"{label}左缩进应为{expected.get('left_indent_pt', 0.0)}磅，实际{actual_left:.2f}磅")
    before = fmt.space_before.pt if fmt.space_before is not None else 0.0
    after = fmt.space_after.pt if fmt.space_after is not None else 0.0
    if not close(before, float(expected.get("space_before_pt", 0.0))):
        errors.append(f"{label}段前应为{expected.get('space_before_pt', 0.0)}磅，实际{before:.2f}磅")
    if not close(after, float(expected.get("space_after_pt", 0.0))):
        errors.append(f"{label}段后应为{expected.get('space_after_pt', 0.0)}磅，实际{after:.2f}磅")
    spacing_kind, spacing_value = line_spacing_value(paragraph)
    expected_spacing = expected["line_spacing"]
    if spacing_kind != expected_spacing["kind"] or spacing_value is None or not close(spacing_value, float(expected_spacing["value"]), 0.08):
        errors.append(f"{label}行距应为{expected_spacing['kind']} {expected_spacing['value']}，实际{spacing_kind} {spacing_value}")
    for run_index, run in enumerate((run for run in paragraph.runs if run.text.strip()), 1):
        run_label = f"{label}第{run_index}个文字片段"
        if east_asia_font(run) != expected["font_east_asia"]:
            errors.append(f"{run_label}中文字体应为{expected['font_east_asia']}，实际{east_asia_font(run)}")
        if run.font.name != expected["font_ascii"]:
            errors.append(f"{run_label}西文字体应为{expected['font_ascii']}，实际{run.font.name}")
        size = run.font.size.pt if run.font.size is not None else None
        if not close(size, float(expected["size_pt"]), 0.1):
            errors.append(f"{run_label}字号应为{expected['size_pt']}磅，实际{size}")
        if bool(run.bold) != bool(expected.get("bold")):
            errors.append(f"{run_label}加粗应为{bool(expected.get('bold'))}，实际{bool(run.bold)}")
    return errors


def audit(
    path: Path,
    material_id: str,
    reference: Path | None = None,
    contracts_path: Path = DEFAULT_CONTRACT_PATH,
) -> tuple[list[str], list[str]]:
    contract = material_contract(material_id, contracts_path)
    if path.suffix.lower() != ".docx":
        return ["字体字号合同审计仅支持DOCX"], []
    document = Document(path)
    if contract.get("mode") == "official-exact":
        if reference is None or not reference.is_file():
            return [f"{material_id}为official-exact，必须提供存在的官方参考模板"], []
        return compare_official_style(reference, path, document), []
    if contract.get("output_format") != "docx":
        return [f"{material_id}不是DOCX版式合同"], []
    errors: list[str] = []
    page = contract["page"]
    expected_size = (297.0, 210.0) if page["orientation"] == "landscape" else (210.0, 297.0)
    for index, section in enumerate(document.sections, 1):
        actual_size = (round(section.page_width.mm, 1), round(section.page_height.mm, 1))
        if any(abs(left - right) > 1.0 for left, right in zip(actual_size, expected_size)):
            errors.append(f"第{index}节纸张方向/尺寸应为{page['orientation']} A4，实际{actual_size}")
        for field, actual in (
            ("top_mm", section.top_margin.mm), ("bottom_mm", section.bottom_margin.mm),
            ("left_mm", section.left_margin.mm), ("right_mm", section.right_margin.mm),
            ("header_mm", section.header_distance.mm), ("footer_mm", section.footer_distance.mm),
        ):
            if abs(actual - float(page[field])) > 1.0:
                errors.append(f"第{index}节{field}应为{page[field]}mm，实际{actual:.1f}mm")
    nonempty_index = 0
    for paragraph_index, paragraph in enumerate(document.paragraphs, 1):
        if not paragraph.text.strip():
            continue
        role = classify(paragraph, nonempty_index, contract)
        errors.extend(check_paragraph(paragraph, role, contract, f"正文第{paragraph_index}段({role})"))
        nonempty_index += 1
        if len(errors) >= 60:
            break
    if len(errors) < 60:
        for table_index, table in enumerate(document.tables, 1):
            errors.extend(check_table_contract(table, contract, f"表{table_index}"))
            for row_index, row in enumerate(table.rows, 1):
                for cell_index, cell in enumerate(row.cells, 1):
                    role = "table_header" if row_index == 1 else "table_body"
                    if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                        errors.append(f"表{table_index}第{row_index}行{cell_index}列必须垂直居中")
                    for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                        if not paragraph.text.strip():
                            continue
                        label = f"表{table_index}第{row_index}行{cell_index}列第{paragraph_index}段({role})"
                        errors.extend(check_paragraph(paragraph, role, contract, label))
                        if len(errors) >= 60:
                            break
    if len(errors) < 60:
        for section_index, section in enumerate(document.sections, 1):
            for container_name, container in (("页眉", section.header), ("页脚", section.footer)):
                for paragraph_index, paragraph in enumerate(container.paragraphs, 1):
                    if not paragraph.text.strip() and not paragraph._p.xpath(".//w:fldSimple|.//w:instrText"):
                        continue
                    errors.extend(
                        check_paragraph(
                            paragraph,
                            "header_footer",
                            contract,
                            f"第{section_index}节{container_name}第{paragraph_index}段(header_footer)",
                        )
                    )
    if len(errors) >= 60:
        errors = errors[:60] + ["版式合同差异超过60处，停止继续列出"]
    return errors, []


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--material-id", required=True)
    parser.add_argument("--reference", type=Path)
    parser.add_argument("--contracts", type=Path, default=DEFAULT_CONTRACT_PATH)
    args = parser.parse_args()
    try:
        errors, warnings = audit(args.docx, args.material_id, args.reference, args.contracts)
    except (OSError, ValueError) as exc:
        print(f"字体字号合同审计失败：{exc}", file=sys.stderr)
        return 2
    for warning in warnings:
        print(f"警告：{warning}")
    for error in errors:
        print(f"错误：{error}")
    if errors:
        print(f"字体字号合同审计未通过：{len(errors)}个错误")
        return 1
    print("字体字号合同审计通过：0个错误")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
