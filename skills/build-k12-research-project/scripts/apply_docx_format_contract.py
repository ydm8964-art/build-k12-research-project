#!/usr/bin/env python3
"""Apply a material's fixed typography and paragraph contract to a non-official DOCX."""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from format_contracts import DEFAULT_CONTRACT_PATH, material_contract

STYLE_NAMES = {
    "cover_title": "K12 Cover Title",
    "title": "K12 Title",
    "subtitle": "K12 Subtitle",
    "heading1": "K12 Heading 1",
    "heading2": "K12 Heading 2",
    "heading3": "K12 Heading 3",
    "heading4": "K12 Heading 4",
    "body": "K12 Body",
    "list": "K12 List",
    "toc_title": "K12 TOC Title",
    "toc_level1": "K12 TOC Level 1",
    "toc_level2": "K12 TOC Level 2",
    "caption": "K12 Caption",
    "table_header": "K12 Table Header",
    "table_body": "K12 Table Body",
    "table_note": "K12 Table Note",
    "photo_placeholder": "K12 Photo Placeholder",
    "header_footer": "K12 Header Footer",
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}


def set_fonts(target, role: dict) -> None:
    target.name = role["font_ascii"]
    target.size = Pt(float(role["size_pt"]))
    target.bold = bool(role.get("bold"))
    target.italic = bool(role.get("italic", False))
    element = target._element if hasattr(target, "_element") else None
    rpr = element.get_or_add_rPr() if element is not None and hasattr(element, "get_or_add_rPr") else None
    if rpr is None and element is not None:
        rpr = element.rPr
    if rpr is not None:
        fonts = rpr.get_or_add_rFonts()
        fonts.set(qn("w:eastAsia"), role["font_east_asia"])
        fonts.set(qn("w:ascii"), role["font_ascii"])
        fonts.set(qn("w:hAnsi"), role["font_ascii"])
        fonts.set(qn("w:cs"), role["font_ascii"])


def set_paragraph_format(target, role: dict) -> None:
    target.alignment = ALIGNMENTS[role["alignment"]]
    target.first_line_indent = Pt(float(role.get("first_line_indent_pt", 0.0)))
    target.left_indent = Pt(float(role.get("left_indent_pt", 0.0)))
    target.right_indent = Pt(float(role.get("right_indent_pt", 0.0)))
    target.space_before = Pt(float(role.get("space_before_pt", 0.0)))
    target.space_after = Pt(float(role.get("space_after_pt", 0.0)))
    spacing = role["line_spacing"]
    target.line_spacing = Pt(float(spacing["value"])) if spacing["kind"] == "fixed" else float(spacing["value"])
    target.keep_with_next = bool(role.get("keep_with_next", False))


def ensure_styles(document: Document, contract: dict) -> None:
    for role_name, role in contract["roles"].items():
        style_name = STYLE_NAMES.get(role_name)
        if not style_name:
            continue
        style = document.styles[style_name] if style_name in document.styles else document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        set_fonts(style.font, role)
        set_paragraph_format(style.paragraph_format, role)


def classify(paragraph, index: int, contract: dict) -> str:
    text = paragraph.text.strip()
    style_name = paragraph.style.name.lower()
    for role, fixed_name in STYLE_NAMES.items():
        if style_name == fixed_name.lower():
            return role
    if text.startswith("【待插入真实照片"):
        return "photo_placeholder"
    if text == "目录":
        return "toc_title"
    if style_name.startswith("toc 2") or "目录 2" in style_name:
        return "toc_level2"
    if style_name.startswith("toc") or "目录 1" in style_name:
        return "toc_level1"
    if re.match(r"^(图|表)\s*\d+", text):
        return "caption"
    if "subtitle" in style_name or "副标题" in style_name:
        return "subtitle"
    if "title" in style_name or "标题" in style_name and "heading" not in style_name:
        return "title"
    if "heading 1" in style_name or re.match(r"^[一二三四五六七八九十]+[、．.]", text):
        return "heading1"
    if "heading 2" in style_name or re.match(r"^[（(][一二三四五六七八九十]+[）)]", text):
        return "heading2"
    if "heading 3" in style_name:
        return "heading3"
    if "heading 4" in style_name or re.match(r"^[（(]\d+[）)]", text):
        return "heading4"
    if re.match(r"^\d+[.、．]", text):
        return contract.get("classification", {}).get("numbered_paragraph_role", "heading3")
    if index == 0:
        return "cover_title" if contract["contract_id"] == "casebook-fixed" else "title"
    if re.match(r"^[•●□☐✓√-]\s*", text):
        return "list"
    return "body"


def apply_paragraph(paragraph, role_name: str, contract: dict) -> None:
    role = contract["roles"][role_name]
    paragraph.style = STYLE_NAMES[role_name]
    set_paragraph_format(paragraph.paragraph_format, role)
    for run in paragraph.runs:
        set_fonts(run.font, role)


def upsert_dxa(parent, tag: str, value: int) -> None:
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    node.set(qn("w:w"), str(value))
    node.set(qn("w:type"), "dxa")


def enforce_table_contract(table, contract: dict) -> None:
    table.autofit = False
    props = table._tbl.tblPr
    layout = props.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        props.append(layout)
    layout.set(qn("w:type"), contract["table"]["layout"])
    upsert_dxa(props, "w:tblInd", 0)

    margins = props.find(qn("w:tblCellMar"))
    if margins is not None:
        props.remove(margins)
    margins = OxmlElement("w:tblCellMar")
    for name, key in (
        ("top", "cell_margin_top_dxa"),
        ("bottom", "cell_margin_bottom_dxa"),
        ("start", "cell_margin_left_dxa"),
        ("end", "cell_margin_right_dxa"),
    ):
        upsert_dxa(margins, f"w:{name}", int(contract["table"][key]))
    props.append(margins)

    if table.rows and contract["table"].get("repeat_header"):
        row_props = table.rows[0]._tr.get_or_add_trPr()
        header = row_props.find(qn("w:tblHeader"))
        if header is None:
            header = OxmlElement("w:tblHeader")
            row_props.append(header)
        header.set(qn("w:val"), "true")


def apply(path: Path, target: Path, material_id: str, contracts_path: Path = DEFAULT_CONTRACT_PATH) -> dict:
    contract = material_contract(material_id, contracts_path)
    if contract.get("mode") == "official-exact":
        raise ValueError(f"{material_id}必须沿用官方模板，禁止用通用合同重排")
    if contract.get("output_format") != "docx":
        raise ValueError(f"{material_id}不是DOCX材料")
    document = Document(path)
    page = contract["page"]
    for section in document.sections:
        if page["orientation"] == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = Mm(297), Mm(210)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = Mm(float(page["top_mm"]))
        section.bottom_margin = Mm(float(page["bottom_mm"]))
        section.left_margin = Mm(float(page["left_mm"]))
        section.right_margin = Mm(float(page["right_mm"]))
        section.header_distance = Mm(float(page["header_mm"]))
        section.footer_distance = Mm(float(page["footer_mm"]))
    ensure_styles(document, contract)
    nonempty_index = 0
    counts: dict[str, int] = {}
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        role = classify(paragraph, nonempty_index, contract)
        apply_paragraph(paragraph, role, contract)
        counts[role] = counts.get(role, 0) + 1
        nonempty_index += 1
    for table in document.tables:
        enforce_table_contract(table, contract)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    role = "table_header" if row_index == 0 else "table_body"
                    apply_paragraph(paragraph, role, contract)
                    counts[role] = counts.get(role, 0) + 1
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                if not paragraph.text.strip() and not paragraph._p.xpath(".//w:fldSimple|.//w:instrText"):
                    continue
                apply_paragraph(paragraph, "header_footer", contract)
                counts["header_footer"] = counts.get("header_footer", 0) + 1
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return {"material_id": material_id, "contract_id": contract["contract_id"], "roles_applied": counts}


def apply_in_place(path: Path, material_id: str, contracts_path: Path = DEFAULT_CONTRACT_PATH) -> dict:
    """Apply a contract through a sibling temporary file, then atomically replace the DOCX."""
    temporary = path.with_name(f".{path.stem}.format-contract.tmp.docx")
    try:
        result = apply(path, temporary, material_id, contracts_path)
        os.replace(temporary, path)
        return result
    finally:
        temporary.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--material-id", required=True)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--contracts", type=Path, default=DEFAULT_CONTRACT_PATH)
    args = parser.parse_args()
    try:
        result = apply(args.docx, args.out, args.material_id, args.contracts)
    except (OSError, ValueError) as exc:
        print(f"应用版式合同失败：{exc}", file=sys.stderr)
        return 1
    print(args.out.resolve())
    print(f"材料{result['material_id']}已应用{result['contract_id']}：{result['roles_applied']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
